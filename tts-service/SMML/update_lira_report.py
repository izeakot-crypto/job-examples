#!/usr/bin/env python3
"""Оновлення Lira_TTS_Звіт.docx — додаємо locale, session_id/comp_schema, кеш"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

doc = Document(r"[USER_HOME]\Desktop\Lira_TTS_Звіт.docx")


def find_paragraph_index(search_text):
    for i, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            return i
    return None


def add_paragraph_after(index, text, style="Normal", bold=False):
    """Додати параграф після вказаного індексу"""
    ref = doc.paragraphs[index]
    new_p = doc.add_paragraph(style=style)
    new_p.text = ""
    run = new_p.add_run(text)
    if bold:
        run.bold = True
    # Move element after reference
    ref._element.addnext(new_p._element)
    return new_p


def add_heading_after(index, text, level=2):
    ref = doc.paragraphs[index]
    new_p = doc.add_heading(text, level=level)
    ref._element.addnext(new_p._element)
    return new_p


# ============================================================
# 1. Оновити секцію 2 — додати session_id/comp_schema
# ============================================================
idx = find_paragraph_index("2.1. Ендпоінти")
if idx:
    # Додаємо опис після заголовка 2.1
    p = add_paragraph_after(idx,
        "Сесія ідентифікується комбінацією session_id + comp_schema (обидва передає клієнт). "
        "session_id — ідентифікатор дзвінка/сесії від LIRA, comp_schema — змінна для біллінгу. "
        "Ключ сесії: session_id_comp_schema.")

# ============================================================
# 2. Оновити секцію 5 — Підтримка мов + locale
# ============================================================
idx = find_paragraph_index("Chirp3-HD підтримує всі 6 цільових мов")
if idx:
    old_text = doc.paragraphs[idx].text
    doc.paragraphs[idx].clear()
    run = doc.paragraphs[idx].add_run(
        "Chirp3-HD підтримує всі 6 цільових мов. "
        "Параметр locale (обов'язковий для не-українських мов) вказує мову синтезу: "
        "uk-UA (за замовчуванням), ru-RU, en-US, pl-PL, es-ES, tr-TR. "
        "Голос (Leda, Puck тощо) працює з усіма мовами — сервер автоматично формує "
        "повне ім'я моделі: {locale}-Chirp3-HD-{voice} (наприклад, ru-RU-Chirp3-HD-Leda)."
    )

# ============================================================
# 3. Додати секцію про кешування — перед Підсумком (секція 10)
# ============================================================
idx_summary = find_paragraph_index("10. Підсумок")
if idx_summary:
    # Змінюємо номери: 10 -> 11
    doc.paragraphs[idx_summary].clear()
    doc.paragraphs[idx_summary].add_run("11. Підсумок").bold = True
    doc.paragraphs[idx_summary].style = doc.styles["Heading 1"]

    # Знаходимо 10.1 і 10.2 і переіменовуємо
    for p in doc.paragraphs:
        if "10.1." in p.text:
            p.clear()
            p.add_run("11.1. Що зроблено").bold = True
        if "10.2." in p.text:
            p.clear()
            p.add_run("11.2. Ключові метрики").bold = True

    # Оновлюємо зміст — додаємо пункт 10
    idx_toc = find_paragraph_index("9. Стрес-тест")
    if idx_toc:
        p_toc = add_paragraph_after(idx_toc, "10. Кешування аудіо", style="List Number")

    # Додаємо повну секцію 10 перед 11. Підсумок
    idx_summary = find_paragraph_index("11. Підсумок")

    # Вставляємо знизу вгору (кожен addnext додає перед Підсумком)
    elements = []

    # Заголовок
    h = doc.add_heading("10. Кешування аудіо", level=1)
    elements.append(h)

    # Опис
    p1 = doc.add_paragraph(
        "Для оптимізації продуктивності та зменшення витрат на Google API реалізовано "
        "дискове кешування згенерованих WAV-файлів."
    )
    elements.append(p1)

    # 10.1
    h2 = doc.add_heading("10.1. Принцип роботи", level=2)
    elements.append(h2)

    bullets = [
        "Ключ кешу = MD5-хеш від комбінації locale + voice + text",
        "Перший запит з унікальним текстом: генерація через Google API (~500-1000мс), результат зберігається на диск",
        "Повторні запити з тим самим текстом, голосом і locale: повертається з кешу (~1-2мс)",
        "TTL (Time To Live): 24 години — файли старші автоматично видаляються фоновим процесом",
        "Різний locale або voice для того ж тексту = різні кеш-записи",
        "Кеш переживає перезапуск сервера (зберігається на диску)",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        elements.append(p)

    # 10.2
    h3 = doc.add_heading("10.2. Результати тестування кешу", level=2)
    elements.append(h3)

    # Таблиця
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Запит", "Cache", "Час"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Перший (генерація)", "MISS", "582мс"],
        ["Повторний (той самий текст)", "HIT", "1мс"],
        ["Той самий текст, інший locale", "MISS", "570мс"],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r+1].cells[c].text = val
    elements.append(table)

    p_result = doc.add_paragraph(
        "Прискорення при cache hit: ~580x (582мс → 1мс). "
        "Для типового колл-центру з однаковими привітаннями та IVR-повідомленнями "
        "очікуваний hit rate: 80-95%."
    )
    elements.append(p_result)

    # 10.3
    h4 = doc.add_heading("10.3. Моніторинг", level=2)
    elements.append(h4)

    p_mon = doc.add_paragraph(
        'Статистика кешу доступна через GET /status → поле "cache": '
        "кількість файлів, розмір, кількість hits/misses, hit rate, TTL. "
        "Заголовок відповіді X-TTS-Cache: HIT або MISS показує джерело для кожного запиту."
    )
    elements.append(p_mon)

    # Вставляємо всі елементи перед Підсумком
    ref = doc.paragraphs[idx_summary]._element
    for elem in reversed(elements):
        ref.addprevious(elem._element if hasattr(elem, '_element') else elem._tbl)

# ============================================================
# 4. Оновити секцію "Що зроблено" — додати locale та кеш
# ============================================================
idx_done = find_paragraph_index("11.1. Що зроблено")
if idx_done:
    # Знаходимо таблицю після цього заголовка
    # Додаємо пункти в список
    last_bullet = None
    for i in range(idx_done + 1, min(idx_done + 20, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if p.style.name == "List Bullet":
            last_bullet = i

    if last_bullet:
        p1 = add_paragraph_after(last_bullet,
            "Параметр locale для вибору мови (uk-UA, ru-RU, en-US, pl-PL, es-ES, tr-TR)",
            style="List Bullet")
        p2 = add_paragraph_after(last_bullet + 1,
            "Дискове кешування аудіо (TTL 24г, hit rate 80-95%)",
            style="List Bullet")
        p3 = add_paragraph_after(last_bullet + 2,
            "Інтеграція з LIRA: session_id + comp_schema від клієнта",
            style="List Bullet")


# Зберігаємо
doc.save(r"[USER_HOME]\Desktop\Lira_TTS_Звіт.docx")
print("OK: Lira_TTS_Звіт.docx updated")
print("  + session_id/comp_schema в секції 2")
print("  + locale в секції 5")
print("  + Нова секція 10: Кешування аудіо")
print("  + Оновлено секцію Підсумок (тепер 11)")

