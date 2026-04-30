import sys, io, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "TTS_Звіт_Окі-Токі.docx")
doc = Document(filepath)

# Завантажуємо результати тесту
results_path = os.path.join(os.path.dirname(__file__), "multilingual_results.json")
with open(results_path, "r", encoding="utf-8") as f:
    data = json.load(f)

# ============================================================
# 1. Замінити таблицю 6 (3.2 Мультимовний тест)
# ============================================================
print("=== Оновлюю таблицю 3.2 мультимовний тест ===")

table6 = doc.tables[6]
table6_element = table6._tbl
parent = table6_element.getparent()

providers = ["Azure", "G.Standard", "G.Wavenet", "Chirp3-HD", "ElevenLabs", "OpenAI"]
langs = ["UA", "EN", "RU", "PL", "ES", "TR"]

header = ["Мова"] + providers
rows = []
for lang in langs:
    row = [lang]
    for prov in providers:
        d = data.get(prov, {}).get(lang)
        if d:
            row.append(f"{d['avg']}с / {d['cps']}")
        else:
            row.append("—")
    rows.append(row)

new_table = doc.add_table(rows=7, cols=7, style='Table Grid')
for ri, row_data in enumerate([header] + rows):
    for ci, cell_text in enumerate(row_data):
        cell = new_table.cell(ri, ci)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

for ci in range(7):
    for run in new_table.cell(0, ci).paragraphs[0].runs:
        run.font.bold = True

table6_element.addprevious(new_table._tbl)
parent.remove(table6_element)

# Оновити підпис
for p in doc.paragraphs:
    txt = p.text.strip()
    if "Формат: час генерації" in txt or "Тестове речення: ~150" in txt or "Усі хмарні API-провайдери" in txt:
        p.text = "Єдиний тест: ~150 символів на мову (привітання колл-центру), 3 спроби, середнє. Формат: час / CPS."
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.italic = True
        print("  Підпис оновлено")
        break

print("  Таблицю 3.2 оновлено з фінальними даними")
print(f"  Провайдери: {', '.join(providers)}")

# Зберігаємо
doc.save(filepath)
print(f"\n✅ Збережено: {filepath}")

# Верифікація
doc2 = Document(filepath)
print("\nТаблиця 3.2:")
for row in doc2.tables[6].rows:
    print(f"  {[c.text[:14] for c in row.cells]}")
