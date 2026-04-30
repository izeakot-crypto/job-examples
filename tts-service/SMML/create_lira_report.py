"""
Звіт по задачі: Інтеграція Google Chirp3-HD TTS — тестовий варіант
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "Lira_TTS_Звіт.docx")
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows, font_size=8):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
            p.paragraph_format.space_after = Pt(0)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                p.paragraph_format.space_after = Pt(0)
    return table

# ============================================================
# ТИТУЛКА
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ЗВІТ ПО ЗАДАЧІ\nІнтеграція Google Chirp3-HD TTS — тестовий варіант')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Lira TTS Server • Платформа Окі-Токі')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Лютий 2026')
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# ЗМІСТ
# ============================================================
doc.add_heading('Зміст', level=1)

toc_items = [
    '1. Задача',
    '2. HTTP-сервер (aiohttp)',
    '3. WebSocket-сервер (стрімінг)',
    '4. Паралельна генерація',
    '5. Підтримка мов',
    '6. Бенчмарки швидкості',
    '7. Залежність часу від довжини повідомлення',
    '8. Мультимовний тест',
    '9. Стрес-тест',
    '10. Підсумок',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 1. ЗАДАЧА
# ============================================================
doc.add_heading('1. Задача', level=1)

doc.add_paragraph('Розгорнути тестовий TTS-сервер на базі Google Chirp3-HD. Сервер приймає HTTP-запит (curl) з текстом, генерує озвучення (WAV 8kHz 16bit mono) і повертає аудіо у відповіді.')

p = doc.add_paragraph()
p.add_run('Підзадачі:').bold = True

tasks = [
    ('HTTP-сервер', 'Python/aiohttp з ендпоінтами /start, /tts, /stop'),
    ('WebSocket-сервер', 'Стрімінг аудіо з мінімальною затримкою'),
    ('Паралельна генерація', 'Розбивка тексту на речення, одночасна генерація через gRPC'),
    ('6 мов', 'UA, EN, RU, PL, ES, TR'),
    ('Бенчмарки', 'Швидкість, мультимовні тести, стрес-тести'),
    ('Залежність від довжини', 'Час генерації vs кількість символів (5–185)'),
    ('Звіти', 'Результати тестування'),
]
for name, desc in tasks:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================
# 2. HTTP-СЕРВЕР
# ============================================================
doc.add_heading('2. HTTP-сервер (aiohttp)', level=1)

doc.add_paragraph('Файл: tts_http_server.py | Порт: 8765 | Хост: 0.0.0.0')

doc.add_heading('2.1. Ендпоінти', level=2)

add_table(doc,
    ['Метод', 'Ендпоінт', 'Призначення', 'Запит (JSON)', 'Відповідь'],
    [
        ['POST', '/start', 'Ініціалізація сесії, warmup gRPC', '{"call_id": "string"}', 'JSON: session_id, voices, startup_ms'],
        ['POST', '/tts', 'Генерація озвучення', '{"session_id": "...", "text": "...", "voice": "Leda"}', 'Binary WAV + HTTP-заголовки'],
        ['POST', '/stop', 'Закриття сесії', '{"session_id": "..."}', 'JSON: status="closed"'],
        ['GET', '/status', 'Статус сервера', '— (без тіла)', 'JSON: active_sessions, config, voices'],
    ],
    font_size=8
)

doc.add_paragraph()

doc.add_heading('2.2. HTTP-заголовки відповіді /tts', level=2)

add_table(doc,
    ['Заголовок', 'Опис', 'Приклад'],
    [
        ['X-TTS-Total-Ms', 'Повний час обробки', '782'],
        ['X-TTS-Gen-Ms', 'Час генерації Google', '680'],
        ['X-TTS-Parts', 'Кількість речень', '3'],
        ['X-TTS-CPS', 'Символів за секунду', '236'],
        ['X-TTS-Audio-Sec', 'Тривалість аудіо', '12.0'],
        ['X-TTS-Text-Len', 'Довжина тексту', '185'],
        ['X-TTS-Voice', 'Голос', 'Leda'],
    ],
    font_size=8
)

doc.add_paragraph()

doc.add_heading('2.3. Конфігурація', level=2)

add_table(doc,
    ['Параметр', 'Значення', 'Опис'],
    [
        ['PORT', '8765', 'Порт сервера'],
        ['MAX_CONCURRENT', '100', 'Макс. одночасних запитів до Google'],
        ['GRPC_CLIENTS', '5', 'Пул gRPC-з\'єднань (HTTP/2 мультиплекс)'],
        ['SILENCE_MS', '150', 'Тиша між реченнями (мс)'],
        ['SESSION_TIMEOUT', '300', 'Таймаут сесії (сек)'],
        ['Формат', 'WAV 8kHz 16bit mono', 'LINEAR16, стандарт телефонії'],
    ],
    font_size=8
)

doc.add_paragraph()

doc.add_heading('2.4. Голоси', level=2)

add_table(doc,
    ['Коротка назва', 'Повна назва Google', 'Стать'],
    [
        ['Leda', 'uk-UA-Chirp3-HD-Leda', 'Жіночий'],
        ['Puck', 'uk-UA-Chirp3-HD-Puck', 'Чоловічий'],
        ['Kore', 'uk-UA-Chirp3-HD-Kore', 'Жіночий'],
        ['Aoede', 'uk-UA-Chirp3-HD-Aoede', 'Жіночий'],
        ['Charon', 'uk-UA-Chirp3-HD-Charon', 'Чоловічий'],
        ['Fenrir', 'uk-UA-Chirp3-HD-Fenrir', 'Чоловічий'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Примітка: ').bold = True
p.add_run('Усього Chirp3-HD має 30 українських голосів (14 жін. + 16 чол.). Для тестового сервера обрано 6 найбільш підходящих для колл-центру.')

# ============================================================
# 3. WEBSOCKET-СЕРВЕР
# ============================================================
doc.add_heading('3. WebSocket-сервер (стрімінг)', level=1)

doc.add_paragraph('Файл: tts_ws_server.py | ws://0.0.0.0:8765')

doc.add_paragraph('WebSocket-версія дозволяє отримувати аудіо по частинах (чанках) — кожне речення надсилається клієнту одразу після генерації, не чекаючи завершення всього тексту.')

doc.add_heading('3.1. Протокол', level=2)

p = doc.add_paragraph()
p.add_run('Клієнт → Сервер (JSON):').bold = True

code = doc.add_paragraph()
run = code.add_run('{"text": "Текст для озвучення", "voice": "Leda", "request_id": "123"}')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Сервер → Клієнт (для кожного чанку):').bold = True

doc.add_paragraph('1) JSON-метадані: type="audio_chunk", part, total_parts, gen_time_ms, wav_bytes')
doc.add_paragraph('2) Бінарні дані: WAV-файл (кожен чанк — повноцінний WAV, можна програвати одразу)')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Фінальне повідомлення:').bold = True
doc.add_paragraph('JSON: type="done", total_ms, parts, voice, text_len, деталі по кожному чанку')

doc.add_heading('3.2. Переваги стрімінгу', level=2)

items = [
    'Перший чанк аудіо приходить через 400-600мс (замість 800-1600мс для повного WAV)',
    'Клієнт починає програвати аудіо ще до завершення генерації всього тексту',
    'Кожен чанк — окремий WAV, не потребує буферизації',
    'MAX_PARALLEL = 3 (паралельна генерація речень)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 4. ПАРАЛЕЛЬНА ГЕНЕРАЦІЯ
# ============================================================
doc.add_heading('4. Паралельна генерація', level=1)

doc.add_paragraph('Ключова оптимізація сервера — розбивка тексту на речення з паралельною генерацією кожного через окремий gRPC-запит.')

doc.add_heading('4.1. Алгоритм', level=2)

steps = [
    'split_sentences(text): regex (?<=[.!?])\\s+ ріже текст на речення',
    'Кожне речення → окремий gRPC-запит через ThreadPoolExecutor',
    'Розподіл по gRPC-клієнтах: client_idx = part_index % 5 (round-robin)',
    'Усі запити виконуються паралельно → час = max(частин), не sum',
    'Склеювання: PCM₁ + 150мс тиші + PCM₂ + ... + PCMₙ → WAV',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('4.2. Ефект оптимізації', level=2)

add_table(doc,
    ['Текст', 'Символів', 'Без split', 'З split (частин)', 'Прискорення'],
    [
        ['1 речення', '~150', '1197мс', '1197мс (1)', '—'],
        ['2 речення', '~150', '—', '580мс (4 ч-ни)', '2×'],
        ['Довгий текст', '~300', '1779мс', '678мс (4 ч-ни)', '2.6×'],
        ['Продакшн', '185', '~1500мс', '782мс (3 ч-ни)', '1.9×'],
    ],
    font_size=9
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Важливо: ').bold = True
p.add_run('Google API має ліміт ~500 символів на один запит. Для довших текстів split обов\'язковий.')

# ============================================================
# 5. ПІДТРИМКА МОВ
# ============================================================
doc.add_heading('5. Підтримка мов', level=1)

doc.add_paragraph('Chirp3-HD підтримує всі 6 цільових мов. Сервер автоматично визначає мову за параметром або використовує uk-UA за замовчуванням.')

add_table(doc,
    ['Мова', 'Код', 'Голос (Leda)', 'Статус'],
    [
        ['Українська', 'uk-UA', 'uk-UA-Chirp3-HD-Leda', '✓ Основна'],
        ['English', 'en-US', 'en-US-Chirp3-HD-Leda', '✓'],
        ['Русский', 'ru-RU', 'ru-RU-Chirp3-HD-Leda', '✓'],
        ['Polski', 'pl-PL', 'pl-PL-Chirp3-HD-Leda', '✓'],
        ['Español', 'es-ES', 'es-ES-Chirp3-HD-Leda', '✓'],
        ['Türkçe', 'tr-TR', 'tr-TR-Chirp3-HD-Leda', '✓'],
    ],
    font_size=9
)

# ============================================================
# 6. БЕНЧМАРКИ
# ============================================================
doc.add_heading('6. Бенчмарки швидкості', level=1)

doc.add_paragraph('Тестове речення (UA, 146 символів): «Дякуємо за дзвінок до компанії Окі-Токі. На жаль, всі оператори зараз зайняті. Будь ласка, залишайтесь на лінії, вам відповідять найближчим часом.»')

doc.add_heading('6.1. Chirp3-HD через сервер (з паралелізацією)', level=2)

add_table(doc,
    ['Метод', 'Час (150 сим)', 'CPS', 'Приріст'],
    [
        ['Базовий (1 запит)', '1.41с', '104', '—'],
        ['gRPC warmup', '1.53с', '95', '+2%'],
        ['Параллель ×2 + warmup', '0.82-1.00с', '146-178', '+36-42%'],
        ['Параллель ×3 + warmup', '0.70-0.87с', '168-209', '+38-44%'],
    ],
    font_size=9
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Стабільний результат (5 запусків, параллель ×2): ').bold = True
p.add_run('Середнє: 0.915с | Мін: 0.862с | Макс: 1.046с | CPS: ~160')

# ============================================================
# 7. ЗАЛЕЖНІСТЬ ВІД ДОВЖИНИ
# ============================================================
doc.add_heading('7. Залежність часу від довжини повідомлення', level=1)

doc.add_paragraph('Тест: 10 фраз від 5 до 185 символів, через сервер з паралельним split. Голос: Chirp3-HD Leda.')

add_table(doc,
    ['#', 'Текст', 'Символів', 'Частин', 'Час (мс)', 'CPS', 'Аудіо (с)'],
    [
        ['1', 'Алло.', '5', '1', '412', '12', '1.32'],
        ['2', 'Дякуємо за дзвінок.', '19', '1', '479', '39', '2.0'],
        ['3', 'Зачекайте, будь ласка.', '22', '1', '430', '51', '2.04'],
        ['4', 'З\'єдную вас з оператором.', '25', '1', '443', '56', '2.0'],
        ['5', 'Натисніть один для з\'єднання з менеджером.', '42', '1', '566', '74', '2.72'],
        ['6', 'Ваш дзвінок дуже важливий для нас, зачекайте.', '45', '1', '620', '72', '3.08'],
        ['7', 'На жаль, усі оператори зайняті. Залишайтесь.', '59', '2', '538', '109', '4.12'],
        ['8', 'Дякуємо за дзвінок до Окі-Токі. Переведено у чергу.', '83', '2', '614', '135', '6.04'],
        ['9', 'Натисніть 1 для оператора. 2 для дзвінка. 0 для повтору.', '108', '3', '587', '183', '7.2'],
        ['10', 'Дякуємо за дзвінок... усі оператори зайняті... залишайтесь...', '185', '3', '782', '236', '12.0'],
    ],
    font_size=8
)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Висновки:').bold = True

items = [
    'Мінімальний час — ~400-430мс (навіть для 5 символів — це overhead gRPC-з\'єднання)',
    'Для коротких фраз (до 50 символів) — 1 частина, 400-620мс',
    'Паралельний split вмикається від 2 речень (~50+ символів) → різке зростання CPS',
    '185 символів з 3 частинами = 782мс (CPS 236) — без split було б ~1500мс',
    'Залежність майже лінійна до ~50 символів, потім виходить на плато завдяки паралелізації',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 8. МУЛЬТИМОВНИЙ ТЕСТ
# ============================================================
doc.add_heading('8. Мультимовний тест', level=1)

doc.add_paragraph('Chirp3-HD Leda, gRPC напряму (без HTTP-сервера), ~150 символів, 3 спроби, найкращий час.')

add_table(doc,
    ['Мова', 'Символів', 'Найкращий', 'Середній', 'Аудіо', 'CPS'],
    [
        ['PL', '152', '984мс', '1239мс', '8.3с', '155'],
        ['EN', '155', '1103мс', '1254мс', '8.4с', '141'],
        ['ES', '161', '1119мс', '1267мс', '9.9с', '144'],
        ['TR', '152', '1157мс', '1182мс', '10.0с', '131'],
        ['UA', '146', '1444мс', '1593мс', '9.8с', '101'],
        ['RU', '148', '1573мс', '1666мс', '9.7с', '94'],
    ],
    font_size=9
)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Висновки:').bold = True

items = [
    'Усі 6 мов працюють без помилок',
    'Найшвидша — PL (984мс), потім EN (1103мс)',
    'Найповільніша — RU (1573мс), потім UA (1444мс)',
    'Латинські мови (PL, EN, ES, TR) на ~30% швидші за кириличні (UA, RU)',
    'Середнє по всіх мовах: 1230мс (best)',
    'З паралельною генерацією (×2) очікуваний час: ~600-800мс',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 9. СТРЕС-ТЕСТ
# ============================================================
doc.add_heading('9. Стрес-тест', level=1)

doc.add_paragraph('Одночасні запити до HTTP-сервера. Текст: «Дякуємо за дзвінок до компанії Окі-Токі. Будь ласка, зачекайте.» (63 символи). Голос: Leda.')

add_table(doc,
    ['Конкурентність', 'Запити', 'OK', 'Помилки', 'Wall (мс)', 'Avg Srv (мс)', 'Min Srv', 'Max Srv'],
    [
        ['1', '1', '1', '0', '—', '509-839', '—', '—'],
        ['5', '5', '5', '0', '—', '537-547', '—', '—'],
        ['10', '10', '10', '0', '—', '481-610', '—', '—'],
        ['20', '20', '20', '0', '—', '458-677', '—', '—'],
        ['50', '50', '50', '0', '4121', '817', '435', '1612'],
        ['70', '70', '70', '0', '5079', '1565', '505', '2819'],
        ['100', '100', '100', '0', '—', '498-531', '—', '—'],
    ],
    font_size=8
)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Висновки:').bold = True

items = [
    '0 помилок на УСІХ рівнях конкурентності (1-100)',
    'При 50 одночасних: avg 817мс — прийнятний час',
    'При 70 одночасних: avg 1565мс, max 2819мс — працює, але повільніше',
    'При 100 одночасних: сервер тримає навантаження',
    'Вузьке місце — Google API, а не наш сервер',
    'Для продакшину рекомендовано до 50 одночасних запитів',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 10. ПІДСУМОК
# ============================================================
doc.add_heading('10. Підсумок', level=1)

doc.add_heading('10.1. Що зроблено', level=2)

add_table(doc,
    ['Задача', 'Статус', 'Результат'],
    [
        ['HTTP-сервер (aiohttp)', '✓ Готово', '/start, /tts, /stop, /status — працюють'],
        ['WebSocket-сервер', '✓ Готово', 'Стрімінг по чанках, кожен = окремий WAV'],
        ['Паралельна генерація', '✓ Готово', 'Split + gRPC pool — прискорення 2-2.6×'],
        ['Підтримка 6 мов', '✓ Готово', 'UA, EN, RU, PL, ES, TR — всі працюють'],
        ['Бенчмарки швидкості', '✓ Готово', '150 сим: 0.9с (парал.), 1.5с (без)'],
        ['Залежність від довжини', '✓ Готово', '10 фраз 5-185 символів, CPS 12-236'],
        ['Мультимовний тест', '✓ Готово', '6 мов, avg 1230мс, 0 помилок'],
        ['Стрес-тест', '✓ Готово', '1-100 concurrent, 0 помилок'],
        ['Звіти', '✓ Готово', 'Цей документ + curl-довідник'],
    ],
    font_size=8
)

doc.add_paragraph()

doc.add_heading('10.2. Ключові метрики', level=2)

add_table(doc,
    ['Метрика', 'Значення'],
    [
        ['Час генерації (150 сим, паралель)', '0.86-1.0с'],
        ['Час генерації (185 сим, 3 частини)', '782мс (CPS 236)'],
        ['Мінімальний час (5 сим)', '412мс'],
        ['Максимальна конкурентність', '100 запитів, 0 помилок'],
        ['Підтримка мов', '6/6 працюють'],
        ['Кількість голосів UA', '30 (14 жін. + 16 чол.)'],
        ['Формат аудіо', 'WAV 8kHz 16bit mono'],
    ],
    font_size=9
)

# ЗБЕРЕЖЕННЯ
doc.save(filepath)
print(f'✅ Збережено: {filepath}')

doc2 = Document(filepath)
print(f'Параграфів: {len(doc2.paragraphs)}')
print(f'Таблиць: {len(doc2.tables)}')
headings = [p.text for p in doc2.paragraphs if p.style and 'Heading' in p.style.name]
for h in headings:
    print(f'  {h}')
