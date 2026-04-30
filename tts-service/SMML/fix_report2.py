import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "TTS_Звіт_Окі-Токі.docx")
doc = Document(filepath)

# Знаходимо всі неправильно вставлені параграфи 2.6-2.9 і видаляємо
to_remove = []
for i, p in enumerate(doc.paragraphs):
    if any(x in p.text for x in [
        "2.6. StyleTTS2", "2.7. XTTS v2", "2.8. Qwen3-TTS", "2.9. Piper TTS",
        "Спеціалізована локальна модель для української мови. 31 голос",
        "Локальна модель з клонуванням голосу. Вимагає GPU. Підтримує EN",
        "Високоякісний TTS з клонуванням. Підтримує тільки 3 мови",
        "Не працює на Windows (помилка espeakbridge)",
    ]):
        to_remove.append(p)

print(f"Видаляю {len(to_remove)} неправильних параграфів...")
for p in to_remove:
    p._element.getparent().remove(p._element)

# Тепер вставляємо правильно — використовуємо addnext від останнього параграфу секції 2
# Знаходимо останній параграф перед секцією 3
idx_sec3 = None
for i, p in enumerate(doc.paragraphs):
    if "3. Результати тестів швидкості" in p.text:
        idx_sec3 = i
        break

print(f"Секція 3 на позиції: {idx_sec3}")

# Знаходимо останній параграф секції 2 (перед секцією 3)
last_sec2 = doc.paragraphs[idx_sec3 - 1]
print(f"Останній параграф секції 2: '{last_sec2.text[:60]}'")

# Вставляємо 2.6-2.9 ПЕРЕД секцією 3, по одному, зверху вниз
ref = doc.paragraphs[idx_sec3]._element  # секція 3

items = [
    ("Heading 2", "2.6. StyleTTS2 Ukrainian (локальний GPU)"),
    ("Normal", "Спеціалізована локальна модель для української мови. 31 голос. Вимагає GPU. Відмінна якість звучання (⭐⭐⭐⭐⭐), але повільна генерація (4.86с, CPS: 10). Success rate: 98.5%. Підтримує тільки українську мову."),
    ("Heading 2", "2.7. XTTS v2 — Coqui TTS (локальний GPU)"),
    ("Normal", "Локальна модель з клонуванням голосу. Вимагає GPU. Підтримує EN, RU, PL, ES, TR — НЕ підтримує українську. Клонування нестабільне (50/50). Середній час: 50с, CPS: 14.5."),
    ("Heading 2", "2.8. Qwen3-TTS — Alibaba (локальний GPU) ❌"),
    ("Normal", "Високоякісний TTS з клонуванням. Підтримує тільки 3 мови (EN, RU, ES). Середній час генерації 171с (CPS: 5) — неприйнятно повільний для колл-центру."),
    ("Heading 2", "2.9. Piper TTS (локальний CPU) ❌"),
    ("Normal", "Не працює на Windows (помилка espeakbridge). Не тестувався."),
]

# Вставляємо в зворотному порядку через addprevious,
# але тепер кожен крок оновлюємо ref щоб зберігти порядок
for style_name, text in reversed(items):
    new_p = doc.add_paragraph(text, style=style_name)
    ref.addprevious(new_p._element)
    ref = new_p._element  # тепер наступний елемент вставиться ПЕРЕД цим

doc.save(filepath)
print(f"\nЗвіт виправлено: {filepath}")

# Перевіримо порядок
doc2 = Document(filepath)
print("\nПеревірка порядку секцій 2.5 - 3:")
printing = False
for p in doc2.paragraphs:
    if "2.5." in p.text:
        printing = True
    if printing:
        if p.text.strip():
            print(f"  [{p.style.name}] {p.text[:80]}")
    if "3.1." in p.text:
        break
