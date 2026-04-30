import sys, io, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "TTS_Звіт_Окі-Токі.docx")
doc = Document(filepath)

base = os.path.dirname(__file__)
with open(os.path.join(base, "multilingual_results.json"), "r", encoding="utf-8") as f:
    data = json.load(f)
with open(os.path.join(base, "edge_results.json"), "r", encoding="utf-8") as f:
    edge = json.load(f)

data["Edge TTS"] = edge

# ============================================================
# Замінити таблицю 6 (3.2 Мультимовний тест) — 7 провайдерів
# ============================================================
print("=== Оновлюю таблицю 3.2 ===")

table6 = doc.tables[6]
table6_element = table6._tbl
parent = table6_element.getparent()

providers = ["Azure", "G.Standard", "G.Wavenet", "Chirp3-HD", "ElevenLabs", "OpenAI", "Edge TTS"]
langs = ["UA", "EN", "RU", "PL", "ES", "TR"]

header = ["Мова"] + providers
rows = []
for lang in langs:
    row = [lang]
    for prov in providers:
        d = data.get(prov, {}).get(lang)
        if d:
            row.append(f"{d['avg']}с / {d['cps']}")
        else:
            row.append("✗ помилка")
    rows.append(row)

new_table = doc.add_table(rows=7, cols=8, style='Table Grid')
for ri, row_data in enumerate([header] + rows):
    for ci, cell_text in enumerate(row_data):
        cell = new_table.cell(ri, ci)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

for ci in range(8):
    for run in new_table.cell(0, ci).paragraphs[0].runs:
        run.font.bold = True

table6_element.addprevious(new_table._tbl)
parent.remove(table6_element)

# Оновити підпис
for p in doc.paragraphs:
    if "Єдиний тест" in p.text or "Формат: час генерації" in p.text or "Тестове речення: ~150" in p.text:
        p.text = "Єдиний тест: одне речення ~150 символів на мову (привітання колл-центру), 3 спроби, середнє. Формат: час / CPS. Edge TTS — нестабільний, PL не працює."
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.italic = True
        break

doc.save(filepath)
print(f"✅ Збережено: {filepath}")

# Верифікація
doc2 = Document(filepath)
print("\nТаблиця 3.2:")
for row in doc2.tables[6].rows:
    print(f"  {[c.text[:14] for c in row.cells]}")
