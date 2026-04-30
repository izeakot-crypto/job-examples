"""
Создание SSML-отчёта для Oki-Toki — полная переработка.
Фокус: эмоции в TTS, SSML-разметка, подходы к генерации эмоциональной речи.
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "SSML_Отчёт_Oki-Toki.docx")

doc = Document()

# ============================================================
# Стили
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# Помощники
def add_table(doc, headers, rows, col_widths=None, font_size=8):
    """Добавить таблицу с форматированием."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовок
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    # Данные
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    # Ширина колонок
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    return table

def add_bold_text(paragraph, bold_text, normal_text):
    """Добавить жирный + обычный текст в один параграф."""
    run = paragraph.add_run(bold_text)
    run.bold = True
    run.font.size = Pt(10)
    run2 = paragraph.add_run(normal_text)
    run2.font.size = Pt(10)

def add_code_block(doc, code_text, font_size=8):
    """Добавить блок кода."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    # Фон через shading
    shading = run._r.get_or_add_tcPr = None  # skip complex shading
    return p

def note(doc, text, italic=True, size=9):
    """Добавить примечание мелким шрифтом."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

# ============================================================
# ТИТУЛЬНАЯ СТРАНИЦА
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ИССЛЕДОВАНИЕ ЭМОЦИОНАЛЬНОГО СИНТЕЗА\nИ SSML-РАЗМЕТКИ ДЛЯ TTS')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Платформа Oki-Toki • Колл-центр')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Проблема: быстрые TTS-модели звучат монотонно.\nЗадача: исследовать способы добавления эмоций в синтез речи.')
run.font.size = Pt(11)
run.font.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Февраль 2026')
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# СОДЕРЖАНИЕ
# ============================================================
doc.add_heading('Содержание', level=1)

toc = [
    '1. Постановка задачи',
    '2. Можно ли добавить эмоции в готовый WAV?',
    '3. SSML-разметка: что это и как работает',
    '    3.1. Что такое SSML',
    '    3.2. Azure Neural TTS — тестирование SSML',
    '    3.3. Google Wavenet — тестирование SSML',
    '    3.4. Сравнение SSML: Azure vs Wavenet',
    '    3.5. ВАЖНО: Unicode-ударения ломают Azure',
    '4. Эмоциональный синтез без SSML: Google Chirp3-HD',
    '    4.1. Как это работает',
    '    4.2. Тест 3-х эмоций: Chirp3-HD vs Wavenet (SSML) vs Azure (SSML)',
    '    4.3. 10 эмоциональных сценариев колл-центра',
    '5. Подходы к генерации SSML-разметки',
    '    5.1. Подход А: Через промт чатбота (GPT генерирует ответ + SSML)',
    '    5.2. Подход Б: Два прохода LLM (текст → SSML → TTS)',
    '    5.3. Подход В: Python-скрипт с ключевыми словами',
    '    5.4. Сравнение всех подходов',
    '6. Замер скорости: с SSML и без',
    '    6.1. Время генерации по эмоциям',
    '    6.2. Варианты пайплайна: время и качество',
    '7. Выводы и рекомендации',
    '    7.1. Ответы на поставленные вопросы',
    '    7.2. Рекомендация для продакшина',
]

for item in toc:
    p = doc.add_paragraph(item, style='List Number' if not item.startswith('    ') else 'Normal')
    if item.startswith('    '):
        p.paragraph_format.left_indent = Cm(1.5)
    for run in p.runs:
        run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# 1. ПОСТАНОВКА ЗАДАЧИ
# ============================================================
doc.add_heading('1. Постановка задачи', level=1)

p = doc.add_paragraph()
p.add_run('Проблема: ').bold = True
p.add_run('У нас есть быстрые TTS-модели (Azure Neural: 0.1-0.5с, Google Wavenet: 0.4-0.8с), но они звучат монотонно — без эмоций, без выразительности. Для колл-центра это критично: клиент слышит робота, а не живого оператора.')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Задачи исследования:').bold = True

items = [
    'Есть ли способы добавить эмоции в уже сгенерированный WAV-файл?',
    'Какая дополнительная разметка (SSML) нужна для добавления эмоций? Какие ключевые слова/параметры требуют модели?',
    'Сделать образцы с несколькими эмоциями в одном сообщении. Измерить генерацию с SSML и без.',
    'Протестировать подходы к созданию SSML: а) через промт чатбота, б) через доп. проход LLM, в) через Python-скрипт с ключевыми словами.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Контекст: ').bold = True
p.add_run('Платформа Oki-Toki — облачный колл-центр. Формат аудио: WAV 8kHz 16bit mono (стандарт телефонии). Основной язык — украинский (+ EN, RU, PL, ES, TR).')

# ============================================================
# 2. МОЖНО ЛИ ДОБАВИТЬ ЭМОЦИИ В ГОТОВЫЙ WAV?
# ============================================================
doc.add_heading('2. Можно ли добавить эмоции в готовый WAV?', level=1)

p = doc.add_paragraph()
run = p.add_run('Короткий ответ: НЕТ.')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph('Эмоции в речи определяются интонацией (F0-контур), ритмом, паузами, тембром — всё это формируется на этапе синтеза. После генерации WAV-файла доступны только примитивные DSP-манипуляции:')

add_table(doc,
    ['Метод', 'Что делает', 'Даёт ли эмоции?'],
    [
        ['Pitch shifting', 'Повышает/понижает тон', 'Нет — звучит как ускоренная запись'],
        ['Time stretching', 'Ускоряет/замедляет', 'Нет — не меняет интонацию'],
        ['EQ / фильтры', 'Меняет частотный баланс', 'Нет — только тембр'],
        ['Реверберация', 'Добавляет «пространство»', 'Нет — не имеет отношения к эмоциям'],
        ['Voice conversion AI', 'Меняет голос (SVC)', 'Частично — но добавляет 2-5с и артефакты'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Вывод: ').bold = True
p.add_run('Эмоции нужно закладывать ДО генерации WAV — через SSML-разметку или через выбор модели, которая сама понимает эмоции из текста.')

# ============================================================
# 3. SSML-РАЗМЕТКА
# ============================================================
doc.add_heading('3. SSML-разметка: что это и как работает', level=1)

doc.add_heading('3.1. Что такое SSML', level=2)

doc.add_paragraph('SSML (Speech Synthesis Markup Language) — XML-стандарт W3C для управления синтезом речи. Основные параметры:')

add_table(doc,
    ['Параметр', 'Что управляет', 'Пример SSML'],
    [
        ['prosody rate', 'Скорость речи', '<prosody rate="85%">медленнее</prosody>'],
        ['prosody pitch', 'Высота тона', '<prosody pitch="+3st">выше</prosody>'],
        ['prosody volume', 'Громкость', '<prosody volume="soft">тише</prosody>'],
        ['break', 'Пауза', '<break time="400ms"/>'],
        ['emphasis', 'Ударение / акцент', '<emphasis level="strong">важно</emphasis>'],
        ['express-as', 'Эмоциональный стиль', '<express-as style="sad">текст</express-as>'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Идея: ').bold = True
p.add_run('Если обернуть текст в SSML с правильными параметрами — TTS-движок должен произнести его с нужной интонацией. Но работает ли это на практике?')

# 3.2. Azure SSML
doc.add_heading('3.2. Azure Neural TTS — тестирование SSML', level=2)

doc.add_paragraph('Голос: uk-UA-PolinaNeural (женский). Протестированы ВСЕ SSML-параметры.')

p = doc.add_paragraph()
p.add_run('Тестовое предложение: ').bold = True
p.add_run('«Добрий день! Дякуємо що зателефонували. Чим можу допомогти?»')

doc.add_paragraph()

add_table(doc,
    ['SSML-параметр', 'Результат', 'Комментарий'],
    [
        ['rate (замедление 50-90%)', '✓ Работает', 'Заметный эффект на скорость речи'],
        ['rate (ускорение 100-150%)', '✗ НЕ работает', 'Звучит одинаково при любом значении >100%'],
        ['pitch (-6st..+6st)', '~ Минимальный', 'Меняет тембр, НЕ даёт эмоций'],
        ['volume (soft..x-loud)', '✗ НЕ работает', 'Никакого эффекта при любых значениях'],
        ['break (100-700ms)', '✓ Работает', 'Единственный надёжный параметр'],
        ['emphasis', '✗ НЕ работает', 'Не поддерживается Neural голосами'],
        ['express-as (style="sad")', '✗ НЕ работает', 'Не поддерживается для uk-UA'],
    ],
    font_size=8
)

doc.add_paragraph()
note(doc, 'Ограничение: express-as работает только для zh-CN и en-US (Jenny, Aria). Для украинского — нет эмоциональных стилей. Подтверждено тестом en-US-JennyNeural vs uk-UA-PolinaNeural — ограничения pitch/volume идентичны для обоих языков.')

p = doc.add_paragraph()
run = p.add_run('Пример SSML-разметки (Azure):')
run.bold = True

code = doc.add_paragraph()
run = code.add_run('''<speak version='1.0' xmlns='http://www.w3.org/2001/10/synthesis'
      xml:lang='uk-UA'>
  <voice name='uk-UA-PolinaNeural'>
    <prosody rate="85%" pitch="-2st">
      На жаль, ми змушені відмовити...
    </prosody>
    <break time="400ms"/>
    <prosody rate="90%">
      Якщо бажаєте, я можу запропонувати альтернативу.
    </prosody>
  </voice>
</speak>''')
run.font.name = 'Consolas'
run.font.size = Pt(8)

# 3.3. Google Wavenet SSML
doc.add_heading('3.3. Google Wavenet — тестирование SSML', level=2)

doc.add_paragraph('Голос: uk-UA-Wavenet-B (женский). Протестированы все SSML-параметры.')

add_table(doc,
    ['SSML-параметр', 'Результат', 'Комментарий'],
    [
        ['rate (замедление)', '✓ Работает', 'x-slow даёт +99% длительности'],
        ['rate (ускорение)', '✓ Работает', 'x-fast даёт -33% длительности'],
        ['pitch', '~ Минимальный', 'Только при экстремальных значениях (±6st)'],
        ['volume', '✗ НЕ работает', 'Никакого эффекта'],
        ['break', '✓ Работает', 'Надёжные паузы'],
        ['emphasis', '~ Минимальный', 'Слабо заметно'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Преимущество Wavenet над Azure: ').bold = True
p.add_run('ускорение rate работает (у Azure — нет). Но pitch и volume — одинаково бесполезны у обоих.')

# 3.4. Сравнение
doc.add_heading('3.4. Сравнение SSML: Azure vs Wavenet', level=2)

add_table(doc,
    ['Параметр', 'Azure Neural (Polina)', 'Google Wavenet-B'],
    [
        ['rate ↓ (замедление)', '✓ Работает', '✓ Работает'],
        ['rate ↑ (ускорение)', '✗ Не работает', '✓ Работает'],
        ['pitch', '~ Минимально', '~ Минимально'],
        ['volume', '✗ Не работает', '✗ Не работает'],
        ['break (паузы)', '✓ Работает', '✓ Работает'],
        ['emphasis', '✗ Не работает', '~ Минимально'],
        ['express-as (эмоции)', '✗ Нет для uk-UA', '✗ Нет'],
        ['Итого полезных', '2 из 7', '3 из 7'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Главный вывод: ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('SSML НЕ может передать эмоции ни на одном провайдере для украинского языка. Можно только замедлить речь и добавить паузы — это не эмоция, а просто скорость.')

# 3.5. Unicode ударения
doc.add_heading('3.5. ВАЖНО: Unicode-ударения ломают Azure', level=2)

p = doc.add_paragraph()
run = p.add_run('Обнаружено: ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('Символ ударения Unicode (U+0301, «◌́») полностью ломает произношение слов в Azure Neural TTS. Слова с ударением вимовляються неправильно — читаються побуквенно или пропускаються.')

doc.add_paragraph('Решение: использовать ТОЛЬКО SSML-тег <phoneme> для ударений, НЕ Unicode-символы. Либо полностью избегать ударений — Azure и так корректно расставляет ударения для украинского.')

# ============================================================
# 4. CHIRP3-HD — ЭМОЦИИ БЕЗ SSML
# ============================================================
doc.add_heading('4. Эмоциональный синтез без SSML: Google Chirp3-HD', level=1)

doc.add_heading('4.1. Как это работает', level=2)

doc.add_paragraph('Google Chirp3-HD (модель 2024-2025) — единственный провайдер с реальным эмоциональным синтезом. Модель САМА распознаёт эмоциональный контекст из текста и генерирует соответствующую интонацию.')

items = [
    'НЕ требует SSML — работает на чистом тексте',
    'Распознаёт пунктуацию: «!» = радость/энергия, «...» = грусть/пауза, «?» = вопрос',
    'Распознаёт лексику: «на жаль», «чудові новини», «вибачте» → соответствующий тон',
    '30 украинских голосов (14 жен. + 16 муж.) — все с эмоциональным синтезом',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# 4.2. Тест 3 эмоций
doc.add_heading('4.2. Тест 3-х эмоций: Chirp3-HD vs Wavenet (SSML) vs Azure (SSML)', level=2)

doc.add_paragraph('Тестовые тексты (~150 символов каждый):')

emotions_texts = [
    ('Happy (радість)', '«Чудові новини! Вашу проблему повністю вирішено! Замовлення вже в дорозі і буде у вас завтра вранці. Дякуємо за терпіння!»'),
    ('Sad (сум)', '«На жаль, ми змушені відмовити... Нам дуже прикро, але повернення коштів неможливе після закінчення гарантійного терміну.»'),
    ('Calm (спокій)', '«Будь ласка, залишайтесь на лінії. Ваш дзвінок дуже важливий для нас. Перший вільний спеціаліст зʼєднається з вами найближчим часом.»'),
]

for emotion, text in emotions_texts:
    p = doc.add_paragraph()
    run = p.add_run(f'{emotion}: ')
    run.bold = True
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Результат прослушивания:').bold = True

add_table(doc,
    ['Эмоция', 'Chirp3-HD Leda\n(без SSML)', 'Wavenet-B\n(без SSML)', 'Wavenet-B\n(с SSML)', 'Azure Polina\n(с SSML)'],
    [
        ['Happy', 'Радостная интонация ✓', 'Монотонно', 'Монотонно', 'Монотонно'],
        ['Sad', 'Сочувственный тон ✓', 'Монотонно', 'Монотонно', 'Монотонно'],
        ['Calm', 'Спокойный, ровный ✓', 'Монотонно', 'Монотонно', 'Монотонно'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Результат: ')
run.bold = True
p.add_run('Только Chirp3-HD чётко передаёт эмоции. SSML-разметка на Wavenet и Azure НЕ даёт слышимого эмоционального эффекта — все три эмоции звучат одинаково монотонно.')

# 4.3. 10 сценариев
doc.add_heading('4.3. 10 эмоциональных сценариев колл-центра', level=2)

doc.add_paragraph('Протестированы 10 реалистичных сценариев колл-центра на 3 голосах Chirp3-HD (Leda, Puck, Kore):')

add_table(doc,
    ['#', 'Сценарий', 'Эмоция', 'Chirp3-HD результат'],
    [
        ['1', 'Приветствие нового клиента', 'Радость', '✓ Приветливая, энергичная'],
        ['2', 'Отказ клиенту', 'Сочувствие', '✓ Мягкий, сочувственный тон'],
        ['3', 'Срочное предупреждение', 'Тревога', '✓ Быстрый, тревожный темп'],
        ['4', 'Ожидание на линии', 'Спокойствие', '✓ Ровный, успокаивающий'],
        ['5', 'Извинение за сбой', 'Эмпатия', '✓ Виноватый, мягкий тон'],
        ['6', 'Проблема решена!', 'Позитив', '✓ Радостный, уверенный'],
        ['7', 'Перевод на оператора', 'Деловой', '✓ Профессиональный, нейтральный'],
        ['8', 'Успокоение злого клиента', 'Деэскалация', '✓ Медленный, успокаивающий'],
        ['9', 'Информация о тарифах', 'Информат.', '✓ Чёткий, размеренный'],
        ['10', 'Тёплое прощание', 'Тепло', '✓ Дружелюбный, тёплый'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Все 10 сценариев — без единого тега SSML. ').bold = True
p.add_run('Chirp3-HD автоматически подбирает интонацию по контексту текста. Для Azure/Wavenet даже с SSML-разметкой все 10 сценариев звучали бы одинаково.')

# ============================================================
# 5. ПОДХОДЫ К ГЕНЕРАЦИИ SSML
# ============================================================
doc.add_heading('5. Подходы к генерации SSML-разметки', level=1)

doc.add_paragraph('Если бы SSML давала эмоции — как её генерировать автоматически? Протестировано 3 подхода.')

# 5.1. Подход А
doc.add_heading('5.1. Подход А: Через промт чатбота', level=2)

p = doc.add_paragraph()
p.add_run('Идея: ').bold = True
p.add_run('GPT-4o-mini генерирует ответ клиенту сразу с SSML-разметкой. Один LLM-запрос для всего.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Промт (system):').bold = True

code = doc.add_paragraph()
run = code.add_run('''Ты — оператор колл-центра Oki-Toki. Отвечай на украинском.
Оформляй ответ в SSML-формате для Azure TTS:
- Радостная новость → rate="108%" pitch="+2st"
- Сочувствие → rate="85%" pitch="-2st" volume="soft"
- Спокойное → rate="95%" pitch="+0st"
Обязательно добавляй <break time="300ms"/> между предложениями.
Оборачивай в <speak><voice name='uk-UA-PolinaNeural'>...</voice></speak>''')
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_paragraph()

add_table(doc,
    ['Параметр', 'Значение'],
    [
        ['Время LLM', '~1.5-2.0с'],
        ['Время TTS', '~0.3с'],
        ['Итого', '~1.8-2.3с'],
        ['Качество SSML', 'Валидный в 80% случаев (2 из 10 ошибки)'],
        ['Эмоциональный эффект', 'Нет — Azure не реагирует на pitch/volume'],
    ],
    font_size=9
)

doc.add_paragraph()
note(doc, 'Проблема: GPT генерирует красивый SSML, но Azure всё равно не передаёт эмоции. Плюс 20% невалидного SSML → нужен fallback на plain text.')

# 5.2. Подход Б
doc.add_heading('5.2. Подход Б: Два прохода LLM (текст → SSML → TTS)', level=2)

p = doc.add_paragraph()
p.add_run('Идея: ').bold = True
p.add_run('Первый GPT генерирует текст ответа. Второй GPT — специализированный SSML-маркапер — размечает этот текст.')

doc.add_paragraph()

add_table(doc,
    ['Этап', 'Время', 'Что делает'],
    [
        ['1. GPT → текст ответа', '~1.5с', 'gpt-4o-mini, temp=0.7, max_tokens=200'],
        ['2. GPT → SSML-разметка', '~2.5-4.0с', 'gpt-4o-mini, temp=0.3, max_tokens=500'],
        ['3. Azure TTS', '~0.3с', 'uk-UA-PolinaNeural, Raw8Khz16Bit'],
        ['ИТОГО', '~4.3-5.8с', 'Из них ~4с — бессмысленная SSML-генерация'],
    ],
    font_size=9
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Вердикт: ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('Худший вариант. +2-4 секунды на SSML-генерацию, которая не даёт эмоций. Двойная стоимость GPT-запросов.')

# 5.3. Подход В
doc.add_heading('5.3. Подход В: Python-скрипт с ключевыми словами', level=2)

p = doc.add_paragraph()
p.add_run('Идея: ').bold = True
p.add_run('Без LLM. Python-скрипт анализирует текст, находит ключевые слова → подставляет шаблонный SSML.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Алгоритм:').bold = True

code = doc.add_paragraph()
run = code.add_run('''# Ключевые слова → эмоция
KEYWORDS = {
    "sad":   ["вибач", "шкода", "на жаль", "прикро", "змушені"],
    "happy": ["чудов", "радий", "вирішено", "вітаємо", "дякуємо"],
    "calm":  ["залишайтесь", "очікуйте", "зачекайте", "хвилинку"],
}
# Шаблон SSML
TEMPLATES = {
    "sad":   '<prosody rate="85%" pitch="-2st">',
    "happy": '<prosody rate="108%" pitch="+2st">',
    "calm":  '<prosody rate="92%">',
    "default": '<prosody rate="95%">',
}''')
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_paragraph()

add_table(doc,
    ['Параметр', 'Значение'],
    [
        ['Время разметки', '~0мс (мгновенно)'],
        ['Время TTS', '~0.3с'],
        ['Итого', '~0.3с (самый быстрый)'],
        ['Точность эмоции', '~70% (простые случаи)'],
        ['Эмоциональный эффект', 'Нет — тот же Azure, те же ограничения'],
    ],
    font_size=9
)

doc.add_paragraph()
note(doc, 'Плюс: мгновенная разметка без GPT-запроса. Минус: эффект тот же — Azure не передаёт эмоции через SSML для украинского.')

# 5.4. Сравнение
doc.add_heading('5.4. Сравнение всех подходов', level=2)

add_table(doc,
    ['Подход', 'Время\nпайплайна', 'Стоимость\nGPT', 'SSML\nкачество', 'Эмоции\nв голосе?', 'Вердикт'],
    [
        ['А. GPT (ответ+SSML)', '~2.0с', '1× GPT', '80% валидн.', '✗ Нет', 'Не даёт результата'],
        ['Б. GPT + GPT-SSML', '~5.0с', '2× GPT', '90% валидн.', '✗ Нет', 'Худший вариант'],
        ['В. Python-скрипт', '~0.3с', '0', '100% валидн.', '✗ Нет', 'Быстро, но бесполезно'],
        ['Chirp3-HD (без SSML)', '~1.7с', '0', 'Не нужен', '✓ ДА', 'РЕКОМЕНДОВАНО'],
        ['Chirp3-HD (парал.)', '~0.9с', '0', 'Не нужен', '✓ ДА', 'Оптимально'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Итог: ')
run.bold = True
run.font.size = Pt(11)
p.add_run('Все три подхода к SSML-генерации бессмысленны — Azure и Wavenet физически не передают эмоции для украинского языка. Chirp3-HD решает задачу без SSML.')

# ============================================================
# 6. ЗАМЕР СКОРОСТИ
# ============================================================
doc.add_heading('6. Замер скорости: с SSML и без', level=1)

doc.add_heading('6.1. Время генерации по эмоциям', level=2)

doc.add_paragraph('Время генерации для 3 эмоций (~150 символов):')

add_table(doc,
    ['Эмоция', 'Azure Polina\n(с SSML)', 'Azure Ostap\n(с SSML)', 'Chirp3-HD Leda\n(без SSML)', 'Chirp3-HD Puck\n(без SSML)'],
    [
        ['Happy', '1.10с', '0.78с', '1.90с', '1.69с'],
        ['Sad', '1.50с', '0.93с', '1.86с', '1.98с'],
        ['Calm', '1.21с', '0.89с', '1.71с', '2.04с'],
        ['Среднее', '1.27с', '0.87с', '1.82с', '1.90с'],
    ],
    font_size=9
)

doc.add_paragraph()
note(doc, 'Azure быстрее (~0.9-1.3с), но голос монотонный. Chirp3-HD медленнее (~1.7-2.0с), но с реальными эмоциями. С параллельной генерацией (×2): Chirp3-HD → ~0.9с.')

# 6.2. Варианты пайплайна
doc.add_heading('6.2. Варианты пайплайна: время и качество', level=2)

doc.add_paragraph('Полный пайплайн: GPT-4o-mini генерирует текст → TTS озвучивает.')

add_table(doc,
    ['#', 'Вариант пайплайна', 'GPT', 'SSML', 'TTS', 'Итого', 'Эмоции?'],
    [
        ['0', 'GPT → GPT-SSML → Azure', '1.5с', '4.0с', '0.3с', '5.8с', '✗'],
        ['1', 'GPT → GPT-SSML (structured) → Azure', '1.5с', '2.5с', '0.3с', '4.3с', '✗'],
        ['2', 'GPT → шаблон SSML → Azure', '1.5с', '0мс', '0.3с', '1.8с', '✗'],
        ['3', 'GPT → Chirp3-HD (без SSML)', '1.5с', '—', '1.5с', '3.0с', '✓'],
        ['4', 'GPT → Chirp3-HD (параллель ×2)', '1.5с', '—', '0.8с', '2.3с', '✓'],
        ['5', 'Кеш фраз → Azure', '0с', '—', '0.3с', '0.3с', '✗'],
        ['6', 'Один GPT (текст+SSML) → Azure', '1.5с', '—', '0.3с', '1.8с', '✗'],
    ],
    font_size=8
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Оптимальный вариант (#4): ').bold = True
p.add_run('GPT текст (~1.5с) → Chirp3-HD параллельная генерация (~0.8с) = 2.3с с эмоциями. Без SSML, без дополнительных GPT-запросов.')

# ============================================================
# 7. ВЫВОДЫ И РЕКОМЕНДАЦИИ
# ============================================================
doc.add_heading('7. Выводы и рекомендации', level=1)

doc.add_heading('7.1. Ответы на поставленные вопросы', level=2)

questions = [
    (
        '1. Можно ли добавить эмоции в готовый WAV?',
        'НЕТ. Эмоции определяются интонацией на этапе синтеза. Постобработка WAV (pitch shifting, time stretching) не создаёт эмоции — только искажает звук.'
    ),
    (
        '2. Какая SSML-разметка нужна для эмоций?',
        'Для украинского языка — НИКАКАЯ. Ни Azure, ни Google Wavenet не передают эмоции через SSML (prosody rate/pitch/volume). Работают только паузы (<break>) и замедление (rate ↓). express-as (эмоциональные стили) не поддерживается для uk-UA.'
    ),
    (
        '3. Образцы с эмоциями + замер времени:',
        'Созданы 10 сценариев × 3 голоса = 30 аудиосэмплов. Только Chirp3-HD различимо передаёт эмоции (без SSML). Время: Chirp3-HD ~1.7с / Azure ~0.9с. С параллелизацией Chirp3: ~0.9с.'
    ),
    (
        '4. Подходы к генерации SSML:',
        'Протестированы 3 подхода (промт GPT, доп. LLM, Python-скрипт). Все бессмысленны: Azure не реагирует на SSML для эмоций. Chirp3-HD не требует SSML вообще — решение задачи без дополнительной разметки.'
    ),
]

for q, a in questions:
    p = doc.add_paragraph()
    run_q = p.add_run(q + '\n')
    run_q.bold = True
    run_q.font.size = Pt(10)
    run_a = p.add_run(a)
    run_a.font.size = Pt(10)
    doc.add_paragraph()

# 7.2. Рекомендация
doc.add_heading('7.2. Рекомендация для продакшина', level=2)

p = doc.add_paragraph()
run = p.add_run('Google Chirp3-HD — единственное решение задачи эмоций.')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph()

items = [
    'Автоматический эмоциональный синтез — из контекста текста, без SSML',
    '30 украинских голосов (14 жен. + 16 муж.)',
    'Не нужен дополнительный GPT-запрос для SSML → экономия 2-4с и денег',
    'С параллельной генерацией: ~0.9с для 150 символов',
    'Рекомендованный голос: uk-UA-Chirp3-HD-Leda (жен.) — приятный, деловой тембр',
    'Стоимость: $16/1M символов (на уровне Azure)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Рекомендованная архитектура:').bold = True

items = [
    'Кеш стандартных фраз — преген. WAV (<0.1с, без эмоций, для типовых ответов)',
    'Chirp3-HD параллельная генерация — для динамических ответов (0.9с, с эмоциями)',
    'Azure Neural как fallback — при недоступности Google API (0.3с, без эмоций)',
    'Полный пайплайн: GPT-4o-mini текст (~1.5с) → Chirp3-HD (~0.9с) = ~2.4с с эмоциями',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('SSML-разметка НЕ рекомендуется ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('— дополнительные 2-4 секунды и стоимость GPT-запроса без какого-либо улучшения эмоциональности голоса.')

# ============================================================
# СОХРАНЕНИЕ
# ============================================================
doc.save(filepath)
print(f'✅ Сохранено: {filepath}')

# Верификация
doc2 = Document(filepath)
print(f'\nСтатистика:')
print(f'  Параграфов: {len(doc2.paragraphs)}')
print(f'  Таблиц: {len(doc2.tables)}')

headings = [(p.style.name, p.text[:80]) for p in doc2.paragraphs if p.style and 'Heading' in p.style.name]
print(f'  Заголовков: {len(headings)}')
for style, text in headings:
    prefix = '    ' if '2' in style else '  '
    print(f'{prefix}{text}')
