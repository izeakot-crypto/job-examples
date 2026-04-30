import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "TTS_Звіт_Окі-Токі.docx")
doc = Document(filepath)

# ═══════════════════════════════════════
# Утиліти
# ═══════════════════════════════════════
def find_paragraph_index(doc, text_contains):
    for i, p in enumerate(doc.paragraphs):
        if text_contains in p.text:
            return i
    return -1

def add_table_after(doc, para_index, headers, rows):
    """Додаємо таблицю після параграфу"""
    table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    # Переміщуємо таблицю після потрібного параграфу
    ref = doc.paragraphs[para_index]._element
    ref.addnext(table._element)
    return table

def replace_table(table, headers, rows):
    """Замінюємо вміст таблиці"""
    # Очистити всі рядки крім першого
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)
    # Оновити заголовок
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        if i < len(hdr):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
    # Додати рядки
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < len(row):
                row[i].text = str(val)
                for p in row[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)

# ═══════════════════════════════════════
# 1. FIX: Таблиця 5 (3.1 Швидкість) — додати CPS + тестове речення + усі моделі
# ═══════════════════════════════════════
print("1. Оновлюю таблицю швидкості (3.1)...")

# Оновлюємо текст перед таблицею 5
idx = find_paragraph_index(doc, "Текст ~150-300 символів")
if idx >= 0:
    doc.paragraphs[idx].text = (
        'Тестове речення (UA, 146 символів): «Дякуємо за дзвінок до компанії Окі-Токі. '
        'На жаль, всі оператори зараз зайняті. Будь ласка, залишайтесь на лінії, '
        'вам відповідять найближчим часом.»'
    )

table5 = doc.tables[5]
replace_table(table5,
    ["Провайдер", "Тип", "Модель/Голос", "~150 сим", "CPS", "Стрімінг 1й чанк"],
    [
        ["Azure (S0 paid)", "API", "Polina Neural", "~0.1-0.3с", "487-1460", "~0.05с"],
        ["Azure (F0 free)", "API", "Polina Neural", "0.1-0.5с", "292-1460", "0.05-0.5с"],
        ["Google Wavenet", "API", "Wavenet-B", "0.4-0.6с", "243-365", "Немає"],
        ["Google Standard", "API", "Standard-B", "0.3-0.4с", "365-487", "Немає"],
        ["Google Chirp3-HD", "API", "Leda/Puck/Kore", "1.4-1.9с", "77-104", "0.13-0.21с"],
        ["Google Chirp3-HD", "API (паралель×2)", "Leda", "0.86-1.0с", "146-170", "—"],
        ["ElevenLabs", "API", "Rachel (multilingual_v2)", "1.94с", "75", "~0.3с"],
        ["OpenAI", "API", "tts-1 nova", "3.82с", "38", "—"],
        ["OpenAI", "API", "tts-1-hd nova", "3.76с", "39", "—"],
        ["Edge TTS", "API (безкошт.)", "Polina", "0.5-2.0с", "~150*", "—"],
        ["StyleTTS2 UA", "Локальн. GPU", "31 голос", "4.86с", "10", "—"],
        ["XTTS v2 (Coqui)", "Локальн. GPU", "Клонування", "50с", "14.5", "—"],
        ["Qwen3-TTS", "Локальн. GPU", "Клонування", "171с", "5", "—"],
        ["Piper TTS", "Локальн. CPU", "—", "—", "—", "Не працює на Windows"],
    ]
)

# ═══════════════════════════════════════
# 2. FIX: Таблиця 6 (3.2 Мультимовний) — додати усі провайдери
# ═══════════════════════════════════════
print("2. Оновлюю мультимовну таблицю (3.2)...")

idx32 = find_paragraph_index(doc, "Azure з WebSocket прогрівом vs Google Wavenet")
if idx32 >= 0:
    doc.paragraphs[idx32].text = (
        "Тестове речення: ~150 символів на кожній мові (привітання колл-центру). "
        "Усі хмарні API-провайдери."
    )

table6 = doc.tables[6]
replace_table(table6,
    ["Мова", "Azure Neural", "Google Wavenet", "Chirp3-HD*", "ElevenLabs", "OpenAI tts-1"],
    [
        ["UA", "0.27с / 541 CPS", "0.75с / 195", "1.52с / 96", "1.94с / 75", "3.82с / 38"],
        ["EN", "0.32с / 447", "0.68с / 210", "—", "1.38с / 104", "1.29с / 111"],
        ["RU", "0.29с / 524", "0.71с / 214", "—", "1.52с / 100", "2.71с / 56"],
        ["PL", "0.35с / 414", "0.82с / 177", "—", "—", "—"],
        ["ES", "0.31с / 468", "0.74с / 196", "—", "—", "—"],
        ["TR", "0.28с / 518", "0.79с / 184", "—", "—", "—"],
    ]
)

# Додаємо примітку після таблиці
idx_note = find_paragraph_index(doc, "Оцінка для платного тарифу")
if idx_note >= 0:
    doc.paragraphs[idx_note].text = (
        "* Chirp3-HD — протестований тільки на UA (основна цільова мова). "
        "ElevenLabs/OpenAI — на UA, EN, RU. CPS = символів/секунду."
    )

# ═══════════════════════════════════════
# 3. FIX: Додати провайдерів 2.6-2.9 (StyleTTS2, XTTS, Qwen3, Piper)
# ═══════════════════════════════════════
print("3. Додаю локальні моделі (StyleTTS2, XTTS v2, Qwen3, Piper)...")

# Знаходимо кінець секції 2 (перед секцією 3)
idx_sec3 = find_paragraph_index(doc, "3. Результати тестів швидкості")
if idx_sec3 >= 0:
    # Додаємо параграфи перед секцією 3
    ref_element = doc.paragraphs[idx_sec3]._element

    texts = [
        ("Heading 2", "2.6. StyleTTS2 Ukrainian (локальний GPU)"),
        ("Normal", "Спеціалізована локальна модель для української мови. 31 голос. Вимагає GPU. "
                   "Відмінна якість звучання (⭐⭐⭐⭐⭐), але повільна генерація (4.86с, CPS: 10). "
                   "Success rate: 98.5%. Підтримує тільки українську мову."),
        ("Heading 2", "2.7. XTTS v2 — Coqui TTS (локальний GPU)"),
        ("Normal", "Локальна модель з клонуванням голосу. Вимагає GPU. Підтримує EN, RU, PL, ES, TR — "
                   "НЕ підтримує українську. Клонування нестабільне (50/50). "
                   "Середній час: 50с, CPS: 14.5."),
        ("Heading 2", "2.8. Qwen3-TTS — Alibaba (локальний GPU) ❌"),
        ("Normal", "Високоякісний TTS з клонуванням. Підтримує тільки 3 мови (EN, RU, ES). "
                   "Середній час генерації 171с (CPS: 5) — неприйнятно повільний для колл-центру."),
        ("Heading 2", "2.9. Piper TTS (локальний CPU) ❌"),
        ("Normal", "Не працює на Windows (помилка espeakbridge). Не тестувався."),
    ]

    for style_name, text in reversed(texts):
        new_p = doc.add_paragraph(text, style=style_name)
        ref_element.addprevious(new_p._element)

# ═══════════════════════════════════════
# 4. FIX: Прибрати 5.1 Розрахунок для Окі-Токі
# ═══════════════════════════════════════
print("4. Видаляю '5.1. Розрахунок для Окі-Токі'...")

# Видаляємо параграф 5.1 і текст
idx_51 = find_paragraph_index(doc, "5.1. Розрахунок для Окі-Токі")
if idx_51 >= 0:
    doc.paragraphs[idx_51]._element.getparent().remove(doc.paragraphs[idx_51]._element)

idx_priklad = find_paragraph_index(doc, "Приклад: 100 000 дзвінків")
if idx_priklad >= 0:
    doc.paragraphs[idx_priklad]._element.getparent().remove(doc.paragraphs[idx_priklad]._element)

# Видаляємо таблицю 11 (розрахунок 20M символів)
table11 = doc.tables[11]
table11._element.getparent().remove(table11._element)

# ═══════════════════════════════════════
# 5. FIX: Таблиця 13 → 12 (Підтримка мов) — додати усі провайдери
# ═══════════════════════════════════════
print("5. Оновлюю таблицю підтримки мов...")

# Після видалення таблиці 11, таблиця 13 стала 12
table_langs = doc.tables[12]
replace_table(table_langs,
    ["Мова", "Azure", "G.Standard", "G.Wavenet", "G.Chirp3-HD", "ElevenLabs", "OpenAI", "Edge", "StyleTTS2", "XTTS v2"],
    [
        ["UA", "Polina, Ostap", "1 жін.", "1 жін.", "30 голосів", "✓ (multilang)", "✓ (multilang)", "Polina, Ostap", "31 голос", "✗"],
        ["EN", "100+", "30+", "30+", "30+", "✓", "✓", "100+", "✗", "✓"],
        ["RU", "3", "5", "5", "30", "✓", "✓", "3", "✗", "✓"],
        ["PL", "2", "5", "5", "30", "✓", "✓", "2", "✗", "✓"],
        ["ES", "10+", "4", "4", "30", "✓", "✓", "10+", "✗", "✓"],
        ["TR", "2", "5", "5", "30", "✓", "✓", "2", "✗", "✓"],
    ]
)

# Оновлюємо заголовок секції 7
idx_sec7 = find_paragraph_index(doc, "7. Підтримка мов")
if idx_sec7 >= 0:
    p = doc.paragraphs[idx_sec7]
    p.clear()
    run = p.add_run("7. Підтримка мов (усі провайдери)")
    run.bold = True

# ═══════════════════════════════════════
# 6. FIX: Таблиця якості — додати усі провайдери
# ═══════════════════════════════════════
print("6. Оновлюю таблицю якості (4.1)...")

table8 = doc.tables[8]
replace_table(table8,
    ["Провайдер", "Тип", "Природність", "Для телефонії", "Примітка"],
    [
        ["Google Chirp3-HD", "API", "★★★★★ (9-10)", "★★★★☆", "Найкраща + автоемоції"],
        ["StyleTTS2 UA", "Локальн. GPU", "★★★★★ (9-10)", "★★★★☆", "Відмінна якість, тільки UA, повільний"],
        ["ElevenLabs", "API", "★★★★☆ (7-8)", "★★★★☆", "Краще ніж OpenAI для UA, дорого"],
        ["Azure Polina", "API", "★★★☆☆ (5-6)", "★★★★☆", "Монотонна, без емоцій, швидка"],
        ["Google Wavenet-B", "API", "★★★☆☆ (7-8)", "★★★☆☆", "Трохи роботизований"],
        ["Google Standard-B", "API", "★★☆☆☆ (4)", "★★☆☆☆", "Помітно роботизований"],
        ["OpenAI tts-1", "API", "★★★☆☆ (5-6)", "★★★☆☆", "Середнє для UA, повільний"],
        ["Edge TTS Polina", "API (безкошт.)", "★★★☆☆ (3-4)", "★★★☆☆", "Погані наголоси, нестабільний"],
        ["XTTS v2", "Локальн. GPU", "★★★☆☆ (5-6)", "★★★☆☆", "Клонування 50/50, без UA"],
    ]
)

# ═══════════════════════════════════════
# 7. FIX: Таблиця SSML — оновити з результатами тестів
# ═══════════════════════════════════════
print("7. Оновлюю таблицю SSML...")

table9 = doc.tables[9]
replace_table(table9,
    ["Функція", "Azure UA", "Google Wavenet UA", "Chirp3-HD", "Примітка"],
    [
        ["prosody rate ↓", "Так (працює)", "Так (працює)", "Не потрібен", "Сповільнення мовлення"],
        ["prosody rate ↑", "НІ (ігнорує)", "Так (працює)", "Не потрібен", "Прискорення мовлення"],
        ["prosody pitch", "Мінімально", "Мінімально", "Не потрібен", "Не змінює емоцію"],
        ["prosody volume", "НІ (ігнорує)", "НІ (ігнорує)", "Не потрібен", "Не працює в жодного"],
        ["break time", "Так", "Так", "Не потрібен", "Паузи між фразами"],
        ["emphasis", "НІ", "Мінімально", "Не потрібен", "Слабкий ефект"],
        ["express-as", "НІ для UK", "—", "Автоматично", "Емоції тільки у Chirp3-HD"],
        ["Емоційний синтез", "НІ", "НІ", "ТАК ✓", "Chirp3 розпізнає з тексту"],
    ]
)

# Оновлюємо текст SSML секції
idx_ssml = find_paragraph_index(doc, "Що працює для української мови")
if idx_ssml >= 0:
    doc.paragraphs[idx_ssml].text = (
        "Результати тестування SSML на Azure та Google Wavenet для української мови. "
        "Chirp3-HD не потребує SSML — автоматично розпізнає емоції з контексту тексту."
    )

# ═══════════════════════════════════════
# ЗБЕРІГАЄМО
# ═══════════════════════════════════════
doc.save(filepath)
print(f"\nЗвіт оновлено: {filepath}")
print("Виправлення:")
print("  ✓ Додані усі 9 TTS-провайдерів + 2 відхилені")
print("  ✓ CPS у таблицях швидкості")
print("  ✓ Тестове речення прописане")
print("  ✓ API vs Локальний GPU позначено")
print("  ✓ Мультимовний тест — усі провайдери")
print("  ✓ Видалено '5.1. Розрахунок для Окі-Токі'")
print("  ✓ Підтримка мов — усі провайдери")
print("  ✓ SSML таблиця оновлена з результатами тестів")
print("  ✓ Якість — усі провайдери з типом (API/GPU)")
