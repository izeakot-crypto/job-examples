import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "TTS_Звіт_Окі-Токі.docx")
doc = Document(filepath)

# ============================================================
# 1. Додати секцію 8: Голоси Chirp3-HD (перед висновками)
# ============================================================
print("=== 1. Додаю секцію 8: Голоси Chirp3-HD ===")

# Знаходимо кінець документа — додаємо нові секції в кінець
body = doc.element.body

# Секція 8
h8 = doc.add_paragraph("8. Голоси Google Chirp3-HD (українська)", style="Heading 1")

p8_intro = doc.add_paragraph(
    "Google Chirp3-HD має 30 українських голосів — 14 жіночих та 16 чоловічих. "
    "Це найбільша колекція якісних українських голосів серед усіх протестованих провайдерів. "
    "Назви голосів — імена зірок та міфологічних персонажів.",
    style="Normal"
)

# Таблиця голосів
voices_f = [
    "Achernar", "Aoede", "Autonoe", "Callirrhoe", "Despina", "Erinome", "Gacrux",
    "Kore", "Laomedeia", "Leda", "Pulcherrima", "Sulafat", "Vindemiatrix", "Zephyr"
]
voices_m = [
    "Achird", "Algenib", "Algieba", "Alnilam", "Charon", "Enceladus", "Fenrir",
    "Iapetus", "Orus", "Puck", "Rasalgethi", "Sadachbia", "Sadaltager", "Schedar",
    "Umbriel", "Zubenelgenubi"
]

voice_table = doc.add_table(rows=3, cols=2, style='Table Grid')
voice_table.cell(0, 0).text = "Жіночі (14)"
voice_table.cell(0, 1).text = "Чоловічі (16)"
voice_table.cell(1, 0).text = ", ".join(voices_f[:7])
voice_table.cell(1, 1).text = ", ".join(voices_m[:8])
voice_table.cell(2, 0).text = ", ".join(voices_f[7:])
voice_table.cell(2, 1).text = ", ".join(voices_m[8:])

for row in voice_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
# Заголовок жирний
for ci in range(2):
    for run in voice_table.cell(0, ci).paragraphs[0].runs:
        run.font.bold = True

h8_2 = doc.add_paragraph("8.1. Швидкість генерації", style="Heading 2")
p8_speed = doc.add_paragraph(
    "Різниця швидкості між голосами мінімальна — ±100мс (1.40-1.77с для 146 символів). "
    "Вибір голосу НЕ впливає на продуктивність системи.",
    style="Normal"
)

h8_3 = doc.add_paragraph("8.2. Різниця між голосами", style="Heading 2")
p8_diff = doc.add_paragraph(
    "Google не надає метаданих про характеристики голосів (тембр, вік, стиль). "
    "Усі 30 голосів мають однакові технічні параметри (sample rate 24000 Hz). "
    "Різниця виключно на слух:",
    style="Normal"
)

diffs = [
    "Тембр — від глибокого низького до високого дзвінкого",
    "Вік звучання — від молодого до зрілого",
    "Характер — від спокійного ділового до емоційного живого",
    "Темп мовлення — незначна різниця у природній швидкості",
    "Інтонація — різна манера наголосів та пауз",
]
for d in diffs:
    bp = doc.add_paragraph(d, style="List Bullet")

p8_rec = doc.add_paragraph(
    "Рекомендація: обирати голос на слух під конкретний сценарій. "
    "Для колл-центру Окі-Токі рекомендовані: Leda (жін.) — приємний, діловий тембр; "
    "Puck (чол.) — чіткий, впевнений голос.",
    style="Normal"
)

print("  Секцію 8 додано (голоси + швидкість + різниця)")


# ============================================================
# 2. Додати секцію 9: Висновки та рекомендації
# ============================================================
print("\n=== 2. Додаю секцію 9: Висновки та рекомендації ===")

h9 = doc.add_paragraph("9. Висновки та рекомендації", style="Heading 1")

# 9.1 Зведена таблиця
h9_1 = doc.add_paragraph("9.1. Зведена порівняльна таблиця", style="Heading 2")

summary_header = ["Провайдер", "Швидкість", "Якість", "Ціна", "Емоції", "Стрімінг", "Вердикт"]
summary_data = [
    ["Google Chirp3-HD", "1.5с (0.9 парал.)", "★★★★★", "$16/1M", "✓ Авто", "✓", "РЕКОМЕНДОВАНО"],
    ["Azure Neural", "0.1-0.3с", "★★★☆☆", "$16/1M", "✗", "✓", "Fallback / кеш"],
    ["Google Wavenet", "0.4-0.6с", "★★★☆☆", "$16/1M", "✗", "✗", "Альтернатива"],
    ["ElevenLabs", "1.9с", "★★★★☆", "~$30/1M", "Частково", "✓", "Дорого"],
    ["OpenAI tts-1", "2.4с", "★★★☆☆", "$15/1M", "✗", "✓", "Повільний"],
    ["Piper TTS", "0.12с", "★☆☆☆☆", "Безкошт.", "✗", "✗", "Якість 3/10"],
    ["Edge TTS", "0.5-2с", "★★★☆☆", "Безкошт.", "✗", "✓", "Без SLA"],
    ["StyleTTS2", "4.9с", "★★★★★", "Безкошт.", "✗", "✗", "Тільки UA, GPU"],
    ["XTTS v2", "50с", "★★★☆☆", "Безкошт.", "✗", "✗", "Без UA"],
    ["Qwen3-TTS", "171с", "—", "Безкошт.", "✗", "✗", "Занадто повільний"],
]

sum_table = doc.add_table(rows=11, cols=7, style='Table Grid')
for ri, row_data in enumerate([summary_header] + summary_data):
    for ci, cell_text in enumerate(row_data):
        cell = sum_table.cell(ri, ci)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

for ci in range(7):
    for run in sum_table.cell(0, ci).paragraphs[0].runs:
        run.font.bold = True
# Перший рядок даних (Chirp3-HD) — виділити
for ci in range(7):
    for run in sum_table.cell(1, ci).paragraphs[0].runs:
        run.font.bold = True

# 9.2 Рекомендація
h9_2 = doc.add_paragraph("9.2. Рекомендований провайдер", style="Heading 2")

rec = doc.add_paragraph(
    "Google Chirp3-HD — оптимальний вибір для Окі-Токі. Причини:",
    style="Normal"
)
for run in rec.runs:
    run.font.bold = True

reasons = [
    "Найвища якість голосу (9-10/10) — живі інтонації, правильна українська вимова",
    "Автоматичний емоційний синтез — розпізнає контекст без SSML чи додаткових налаштувань",
    "30 українських голосів — найбільший вибір серед усіх провайдерів",
    "Підтримка 6 мов (UA, EN, RU, PL, ES, TR) — мультимовний колл-центр",
    "Стрімінг (gRPC) — перший чанк за 130-190мс",
    "Паралельна генерація — оптимізація до 0.9с для 150 символів",
    "Ціна $16/1M символів — на рівні Azure, вдвічі дешевше за ElevenLabs",
]
for r in reasons:
    doc.add_paragraph(r, style="List Bullet")

# 9.3 Архітектура
h9_3 = doc.add_paragraph("9.3. Рекомендована архітектура", style="Heading 2")
p_arch = doc.add_paragraph(
    "Гібридна схема для мінімальної затримки:",
    style="Normal"
)

arch_items = [
    "Кеш стандартних фраз — пре-генеровані WAV файли для типових відповідей (<0.1с)",
    "Chirp3-HD паралельна генерація — для динамічних відповідей (0.9с)",
    "Azure Neural як fallback — при недоступності Google API (0.3с)",
    "Повний пайплайн: GPT-4o-mini текст (~1.5с) → Chirp3-HD аудіо (~0.9с) = ~2.4с загалом",
]
for a in arch_items:
    doc.add_paragraph(a, style="List Bullet")

# 9.4 Що НЕ підходить
h9_4 = doc.add_paragraph("9.4. Що не підходить для колл-центру", style="Heading 2")

not_suitable = [
    "SSML — не дає покращення якості. Жоден провайдер не реалізує емоційний синтез через SSML для української",
    "Piper TTS — швидкий, але якість 3/10 (роботний звук)",
    "OpenAI TTS — повільний для UA (2.4с), немає 8kHz WAV",
    "XTTS v2 — не підтримує українську мову",
    "Qwen3-TTS — 171с на генерацію, неприйнятно повільний",
    "Локальні GPU-моделі — потребують дорогого обладнання, складне обслуговування",
]
for n in not_suitable:
    doc.add_paragraph(n, style="List Bullet")

print("  Секцію 9 додано (зведена таблиця + рекомендація + архітектура)")


# ============================================================
# ЗБЕРІГАЄМО
# ============================================================
doc.save(filepath)
print(f"\n✅ Збережено: {filepath}")

# Верифікація
doc2 = Document(filepath)
print(f"\nСекції Heading 1:")
for p in doc2.paragraphs:
    if p.style.name == "Heading 1":
        print(f"  {p.text}")
print(f"Таблиць: {len(doc2.tables)}")
