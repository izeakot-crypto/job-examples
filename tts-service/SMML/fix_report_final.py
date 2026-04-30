import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "TTS_Звіт_Окі-Токі.docx")
doc = Document(filepath)

fixes = []

# ============================================================
# FIX 1: Секція 3 — додати номер "3."
# ============================================================
for i, p in enumerate(doc.paragraphs):
    if p.style and p.style.name == 'Heading 1' and p.text.strip() == 'Тести швидкості':
        p.text = '3. Тести швидкості'
        for run in p.runs:
            run.font.size = None  # inherit from style
        fixes.append(f'FIX 1: [{i}] "Тести швидкості" → "3. Тести швидкості"')
        break

# ============================================================
# FIX 2: Таблиця 5 — видалити порожній останній рядок
# ============================================================
table5 = doc.tables[5]
last_row = table5.rows[-1]
last_cells = [c.text.strip() for c in last_row.cells]
if all(c == '' for c in last_cells):
    tbl = table5._tbl
    tr = last_row._tr
    tbl.remove(tr)
    fixes.append(f'FIX 2: Таблиця 5 — видалено порожній рядок (було {last_cells})')

# ============================================================
# FIX 3: Додати секцію 9.4 "Що не працює / не рекомендується"
# ============================================================
# Знаходимо останній елемент секції 9.3
last_elem = None
found_93 = False
for i, p in enumerate(doc.paragraphs):
    if '9.3' in p.text and 'Heading' in (p.style.name or ''):
        found_93 = True
    if found_93:
        last_elem = p

if last_elem:
    # Додаємо після останнього елементу документа
    # Створюємо 9.4 заголовок
    p94 = doc.add_heading('9.4. Що не працює / не рекомендується', level=2)

    items = [
        ("Qwen3-TTS", "171с на генерацію — абсолютно непридатний для real-time. CPS: 5."),
        ("XTTS v2", "Не підтримує українську мову. Час генерації 50с — занадто повільний."),
        ("StyleTTS2", "Відмінна якість, але 4.86с генерації та вимога GPU роблять його непрактичним для хмарного сервісу."),
        ("Piper TTS", "Найшвидший (CPS: 517), але якість 3/10 — роботний голос, не підходить для колл-центру."),
        ("Edge TTS", "Безкоштовний, але неофіційний API без SLA. PL не працює. Нестабільний (0.5-5с розкид)."),
        ("OpenAI TTS", "Повільний (2.2-3.1с), немає нативного 8kHz, дорогий ($15-30/1M). Не конкурентний."),
    ]

    for name, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run_bold = p.add_run(f'{name}: ')
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_normal = p.add_run(desc)
        run_normal.font.size = Pt(10)

    fixes.append('FIX 3: Додано секцію 9.4 "Що не працює / не рекомендується" (6 пунктів)')

# ============================================================
# FIX 4: Chirp3-HD безкоштовний тариф — уточнити
# ============================================================
# В таблиці цін Chirp3-HD має "—" для безкоштовного, але Google Cloud дає free tier
table10 = doc.tables[10]
for row in table10.rows:
    cells = row.cells
    if 'Chirp3-HD' in cells[0].text:
        if cells[2].text.strip() == '—':
            cells[2].text = '1M сим/міс*'
            for run in cells[2].paragraphs[0].runs:
                run.font.size = Pt(8)
            fixes.append('FIX 4: Chirp3-HD безкоштовний тариф "—" → "1M сим/міс*"')

# ============================================================
# Зберігаємо
# ============================================================
doc.save(filepath)
print(f'✅ Збережено: {filepath}')
print(f'\nВиправлення ({len(fixes)}):')
for f in fixes:
    print(f'  {f}')

# Верифікація
doc2 = Document(filepath)
print(f'\nВерифікація:')

# H1 без номера
for p in doc2.paragraphs:
    if p.style and p.style.name == 'Heading 1':
        if not p.text.strip()[0].isdigit():
            print(f'  ⚠ H1 без номера: "{p.text}"')

# Таблиця 5 — останній рядок
t5 = doc2.tables[5]
last = [c.text.strip() for c in t5.rows[-1].cells]
print(f'  Таблиця 5 останній рядок: {last}')

# 9.4 існує?
for p in doc2.paragraphs:
    if '9.4' in p.text:
        print(f'  Секція 9.4: {p.text[:80]}')
        break

print(f'\nВсього параграфів: {len(doc2.paragraphs)}')
print(f'Всього таблиць: {len(doc2.tables)}')
