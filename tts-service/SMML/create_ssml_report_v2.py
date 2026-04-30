import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_bold_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

# ═══════════════════════════════════════
# ТИТУЛЬНА
# ═══════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("АНАЛИТИЧЕСКИЙ ОТЧЁТ")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Исследование TTS-провайдеров и SSML-разметки\nдля платформы Oki-Toki")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Формат: WAV 8kHz 16bit mono (телефония)\n"
    "6 языков • 9 TTS-провайдеров • 30+ голосов\n"
    "SSML-разметка • Эмоциональный синтез • Оптимизация скорости\n"
    "Облачные API vs Локальные GPU-модели"
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Февраль 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════
# СОДЕРЖАНИЕ
# ═══════════════════════════════════════
add_heading("Содержание", 1)
sections = [
    "1. Цели и задачи исследования",
    "2. Протестированные TTS-провайдеры (9 систем)",
    "3. Сравнение скорости генерации",
    "4. Мультиязычное тестирование (6 языков)",
    "5. SSML-разметка: тестирование и выводы",
    "6. Эмоциональный синтез: сравнение подходов",
    "7. Качество звука: субъективная оценка",
    "8. Оптимизация скорости: методы и результаты",
    "9. Стоимость провайдеров",
    "10. Рекомендуемая архитектура для продакшина",
    "11. Итоговые выводы и рекомендации",
]
for s in sections:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ═══════════════════════════════════════
# 1. ЦЕЛИ
# ═══════════════════════════════════════
add_heading("1. Цели и задачи исследования", 1)
doc.add_paragraph(
    "Цель исследования — определить оптимальный TTS-провайдер для платформы Oki-Toki "
    "(облачный колл-центр) с учётом следующих требований:"
)
requirements = [
    "Формат аудио: WAV 8kHz 16bit mono (стандарт телефонии)",
    "Скорость генерации: менее 1 секунды для фраз до 150 символов",
    "Качество голоса: естественное звучание, максимально близкое к живому оператору",
    "Эмоциональный синтез: способность передавать эмоции (радость, сочувствие, спокойствие)",
    "Мультиязычность: поддержка UA, EN, RU, PL, ES, TR",
    "SSML-разметка: возможность управления просодией (скорость, тон, громкость, паузы)",
    "Стоимость: оптимальное соотношение цена/качество при высоких объёмах",
]
for r in requirements:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# 2. ПРОВАЙДЕРЫ
# ═══════════════════════════════════════
add_heading("2. Протестированные TTS-провайдеры (9 систем)", 1)

doc.add_paragraph(
    "Всего протестировано 9 TTS-систем: 5 облачных (через API) и 4 локальных (требуют GPU/CPU)."
)

add_heading("2.1. Сводная таблица всех TTS", 2)
add_table(
    ["#", "TTS-система", "Тип", "Мови (UA)", "Голосов UA", "Эмоции", "SSML"],
    [
        ["1", "Azure Neural", "Облако (API)", "✓", "2", "Нет", "Ограничен"],
        ["2", "Google Standard", "Облако (API)", "✓", "1", "Нет", "Да"],
        ["3", "Google Wavenet", "Облако (API)", "✓", "1", "Нет", "Частично"],
        ["4", "Google Chirp3-HD", "Облако (API)", "✓", "30", "Автоматич.", "Не нужен"],
        ["5", "ElevenLabs", "Облако (API)", "✓", "21 (мультиязыч.)", "Частично", "Нет"],
        ["6", "OpenAI TTS", "Облако (API)", "✓", "6 (мультиязыч.)", "Нет", "Нет"],
        ["7", "StyleTTS2 UA", "Локальный (GPU)", "✓", "31", "Нет", "Нет"],
        ["8", "XTTS v2 (Coqui)", "Локальный (GPU)", "✗", "Клонирование", "Нет", "Нет"],
        ["9", "Edge TTS", "Облако (бесплатн.)", "✓", "2", "Нет", "Ограничен"],
    ]
)

doc.add_paragraph()
add_note("Также рассматривались, но не прошли тестирование: Qwen3-TTS (только 3 языка, 171с на генерацию — неприемлемо медленный), Piper TTS (не работает на Windows).")

add_heading("2.2. Облачные API-провайдеры", 2)

add_heading("Microsoft Azure Neural TTS", 3)
doc.add_paragraph(
    "Голоса: uk-UA-PolinaNeural (жен.) и uk-UA-OstapNeural (муж.). "
    "Протокол: WebSocket с pre-connect и warmup. Формат: Raw8Khz16BitMonoPcm. "
    "Самый быстрый среди всех облачных провайдеров (0.1-0.3с на S0 тарифе)."
)

add_heading("Google Cloud TTS — 3 модели", 3)
doc.add_paragraph("Standard-B — базовая модель, 1 женский голос, роботизированное звучание.")
doc.add_paragraph("Wavenet-B — нейросетевая модель, 1 женский голос, среднее качество, поддерживает SSML rate.")
doc.add_paragraph(
    "Chirp3-HD — новейшая модель (2024-2025), 30 украинских голосов (14 жен. + 16 муж.). "
    "Автоматическое распознавание эмоций из текста без SSML. Лучшее качество звучания."
)

add_heading("ElevenLabs Multilingual v2", 3)
doc.add_paragraph(
    "Облачный API с 21 мультиязычным голосом. Поддерживает украинский. "
    "Качество звучания хорошее, лучше чем OpenAI для украинского языка. "
    "Высокая стоимость ($99/мес за 100K символов)."
)

add_heading("OpenAI TTS (tts-1 / tts-1-hd)", 3)
doc.add_paragraph(
    "Облачный API с 6 голосами. Поддерживает украинский, но звучание менее естественное. "
    "Модели: tts-1 (быстрая) и tts-1-hd (качественная). "
    "Медленная генерация для украинского языка (3.8с)."
)

add_heading("2.3. Локальные модели (GPU)", 2)

add_heading("StyleTTS2 Ukrainian", 3)
doc.add_paragraph(
    "Специализированная локальная модель для украинского языка. 31 голос. "
    "Требует GPU. Отличное качество звучания (⭐⭐⭐⭐⭐). "
    "Среднее время генерации: 4.86с, CPS: 9.97. Success rate: 98.5%."
)

add_heading("XTTS v2 (Coqui TTS)", 3)
doc.add_paragraph(
    "Локальная модель с клонированием голоса. Требует GPU. "
    "Поддерживает EN, RU, PL, ES, TR — НЕ поддерживает украинский. "
    "Клонирование нестабильно (50/50). Среднее время: 50с."
)

add_heading("Edge TTS (Microsoft)", 3)
doc.add_paragraph(
    "Бесплатный облачный TTS от Microsoft. Поддерживает все 6 языков. "
    "Очень быстрый (CPS ~150), но роботизированное звучание (⭐⭐). "
    "Плохие ударения в украинском языке."
)

add_heading("2.4. Отклонённые модели", 2)
add_table(
    ["Модель", "Причина отклонения"],
    [
        ["Qwen3-TTS (Alibaba)", "Только 3 языка (EN,RU,ES). Генерация 171с — неприемлемо"],
        ["Piper TTS", "Не работает на Windows (ошибка espeakbridge)"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 3. СКОРОСТЬ
# ═══════════════════════════════════════
add_heading("3. Сравнение скорости генерации", 1)

add_note(
    "Тестовое предложение (UA, 146 символов): «Дякуємо за дзвінок до компанії Окі-Токі. "
    "На жаль, всі оператори зараз зайняті. Будь ласка, залишайтесь на лінії, "
    "вам відповідять найближчим часом.»"
)
doc.add_paragraph()

add_heading("3.1. Все провайдеры — украинский язык (146 символов)", 2)
add_table(
    ["#", "TTS-система", "Тип", "Время", "CPS", "Комментарий"],
    [
        ["1", "Azure Polina (warmup)", "API", "0.27с", "541", "Самый быстрый (F0 нестабильный)"],
        ["2", "Google Wavenet-B", "API", "0.75с", "195", "Стабильный"],
        ["3", "Google Standard-B", "API", "0.98с", "149", "Роботизированный"],
        ["4", "OpenAI tts-1", "API", "1.29с*", "113*", "*EN; UA = 3.82с / 38 CPS"],
        ["5", "Azure Ostap", "API", "1.29с", "113", "Нестабильный (F0 tier)"],
        ["6", "ElevenLabs v2", "API", "1.94с", "75", "Хорошее качество"],
        ["7", "Google Chirp3-HD Leda", "API", "1.52с", "96", "Лучшее качество + эмоции"],
        ["8", "OpenAI tts-1-hd", "API", "3.76с", "39", "Медленный для UA"],
        ["9", "StyleTTS2 UA", "Локальн. GPU", "4.86с", "10", "Отличное качество, но медленный"],
        ["10", "Edge TTS", "API (бесплатн.)", "5.50с*", "~150*", "*Данные из отдельного теста"],
        ["11", "XTTS v2", "Локальн. GPU", "50с", "14.5", "Не поддерживает UA"],
        ["12", "Qwen3-TTS", "Локальн. GPU", "171с", "5", "Неприемлемо медленный"],
    ]
)

doc.add_paragraph()
add_note(
    "CPS = Characters Per Second (символов в секунду). Чем больше — тем быстрее. "
    "Тип: API = облачный сервер, Локальн. GPU = требует видеокарту."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 4. МУЛЬТИЯЗЫЧНЫЙ
# ═══════════════════════════════════════
add_heading("4. Мультиязычное тестирование (6 языков)", 1)

add_note(
    "Тестовые предложения: ~150 символов на каждом языке, одинаковый смысл "
    "(приветствие колл-центра). Все результаты — время генерации / CPS."
)
doc.add_paragraph()

add_heading("4.1. Результаты по языкам", 2)
add_table(
    ["Язык", "Azure Neural", "Google Wavenet", "Google Chirp3-HD", "ElevenLabs v2", "OpenAI tts-1"],
    [
        ["UA", "0.27с / 541", "0.75с / 195", "1.52с / 96", "1.94с / 75", "3.82с / 38"],
        ["EN", "0.32с / 447", "0.68с / 210", "—*", "1.38с / 104", "1.29с / 111"],
        ["RU", "0.29с / 524", "0.71с / 214", "—*", "1.52с / 100", "2.71с / 56"],
        ["PL", "0.35с / 414", "0.82с / 177", "—*", "—**", "—**"],
        ["ES", "0.31с / 468", "0.74с / 196", "—*", "—**", "—**"],
        ["TR", "0.28с / 518", "0.79с / 184", "—*", "—**", "—**"],
    ]
)
doc.add_paragraph()
add_note("* Chirp3-HD — протестирован только на UA (основной целевой язык)")
add_note("** ElevenLabs / OpenAI — протестированы на UA, EN, RU (3 из 6 языков)")

add_heading("4.2. Поддержка языков", 2)
add_table(
    ["TTS-система", "UA", "EN", "RU", "PL", "ES", "TR", "Итого"],
    [
        ["Azure Neural", "✓", "✓", "✓", "✓", "✓", "✓", "6/6"],
        ["Google Standard", "✓", "✓", "✓", "✓", "✓", "✓", "6/6"],
        ["Google Wavenet", "✓", "✓", "✓", "✓", "✓", "✓", "6/6"],
        ["Google Chirp3-HD", "✓", "✓", "✓", "✓", "✓", "✓", "6/6"],
        ["ElevenLabs v2", "✓", "✓", "✓", "✓", "✓", "✓", "6/6"],
        ["OpenAI TTS", "✓", "✓", "✓", "✓", "✓", "✓", "6/6"],
        ["Edge TTS", "✓", "✓", "✓", "✓", "✓", "✓", "6/6"],
        ["StyleTTS2 UA", "✓", "✗", "✗", "✗", "✗", "✗", "1/6"],
        ["XTTS v2", "✗", "✓", "✓", "✓", "✓", "✓", "5/6"],
        ["Qwen3-TTS", "✗", "✓", "✓", "✗", "✓", "✗", "3/6"],
        ["Piper TTS", "—", "—", "—", "—", "—", "—", "Не работает"],
    ]
)

doc.add_paragraph()
add_bold_text("Вывод:")
doc.add_paragraph(
    "Все облачные провайдеры (Azure, Google, ElevenLabs, OpenAI, Edge) поддерживают все 6 языков. "
    "Локальные модели имеют ограниченную языковую поддержку: StyleTTS2 — только UA, "
    "XTTS v2 — все кроме UA, Qwen3 — только 3 языка."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 5. SSML
# ═══════════════════════════════════════
add_heading("5. SSML-разметка: тестирование и выводы", 1)

doc.add_paragraph(
    "SSML (Speech Synthesis Markup Language) — стандарт разметки для управления синтезом речи. "
    "Протестировано на Azure Neural и Google Wavenet (Chirp3-HD не поддерживает и не требует SSML)."
)

add_heading("5.1. Azure Neural — SSML-поддержка", 2)
add_table(
    ["SSML-параметр", "Результат", "Комментарий"],
    [
        ["rate (замедление)", "Работает", "50-90% — заметный эффект"],
        ["rate (ускорение)", "Не работает", "100-150% — звучит одинаково"],
        ["pitch", "Минимальный эффект", "Меняет тембр, не эмоцию"],
        ["volume", "Не работает", "Никакого эффекта"],
        ["break (паузы)", "Работает", "Единственный надёжный параметр"],
        ["emphasis", "Не работает", "Не поддерживается для Neural"],
        ["express-as", "Не работает", "Не поддерживается для uk-UA"],
    ]
)

doc.add_paragraph()
add_note(
    "Ограничения SSML одинаковы для ВСЕХ языков Azure Neural. "
    "Подтверждено тестом en-US-JennyNeural vs uk-UA-PolinaNeural — идентичные результаты."
)

add_heading("5.2. Google Wavenet — SSML-поддержка", 2)
add_table(
    ["SSML-параметр", "Результат", "Комментарий"],
    [
        ["rate (замедление)", "Работает", "x-slow даёт +99% длительности"],
        ["rate (ускорение)", "Работает", "x-fast даёт -33% длительности"],
        ["pitch", "Минимальный эффект", "Только при экстремальных значениях"],
        ["volume", "Не работает", "Никакого эффекта"],
        ["break (паузы)", "Работает", "Надёжный параметр"],
        ["emphasis", "Минимальный эффект", "Слабо заметно"],
    ]
)

add_heading("5.3. Chirp3-HD, ElevenLabs, OpenAI, локальные", 2)
doc.add_paragraph(
    "Эти системы НЕ поддерживают SSML. Chirp3-HD не нуждается в SSML — "
    "автоматически распознаёт эмоции из текста. Остальные — без управления просодией."
)

add_heading("5.4. SSML + LLM: тестирование пайплайна", 2)
doc.add_paragraph(
    "Протестирован пайплайн: Диалог → GPT-4o-mini генерирует SSML → Azure TTS."
)
doc.add_paragraph("Общее время: ~5.8с (GPT текст ~1.5с + GPT SSML ~4.0с + TTS ~0.3с)", style='List Bullet')
doc.add_paragraph("2 из 10 тестов — невалидный SSML (fallback на plain text)", style='List Bullet')
doc.add_paragraph("Результат: SSML не давала заметного эмоционального эффекта", style='List Bullet')

doc.add_paragraph()
add_bold_text("Вывод по SSML:")
doc.add_paragraph(
    "SSML-разметка через LLM — дополнительные 2-4 секунды без результата. "
    "Ни один провайдер не способен передать эмоции через SSML. "
    "Chirp3-HD решает задачу эмоционального синтеза без SSML."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 6. ЭМОЦИИ
# ═══════════════════════════════════════
add_heading("6. Эмоциональный синтез: сравнение подходов", 1)

doc.add_paragraph(
    "Тест на 3 эмоциях (радость, грусть, спокойствие) с текстами ~150 символов."
)

add_heading("6.1. Chirp3-HD vs Wavenet vs Azure", 2)
add_table(
    ["Эмоция", "Chirp3-HD Leda", "Wavenet (без SSML)", "Wavenet (с SSML)", "Azure Polina"],
    [
        ["Happy", "Радостная ✓", "Монотонно", "Монотонно", "Монотонно"],
        ["Sad", "Сочувственный ✓", "Монотонно", "Монотонно", "Монотонно"],
        ["Calm", "Спокойный ✓", "Монотонно", "Монотонно", "Монотонно"],
    ]
)

add_heading("6.2. Время генерации по эмоциям", 2)
add_table(
    ["Эмоция", "Azure Polina", "Azure Ostap", "Chirp3 Leda", "Chirp3 Puck"],
    [
        ["Happy", "1.10с", "0.78с", "1.90с", "1.69с"],
        ["Sad", "1.50с", "0.93с", "1.86с", "1.98с"],
        ["Calm", "1.21с", "0.89с", "1.71с", "2.04с"],
    ]
)

add_heading("6.3. Все 30 голосов Chirp3-HD", 2)
doc.add_paragraph(
    "Протестированы все 30 украинских голосов Chirp3-HD (146 символов). "
    "Средний: 1.56с. Разброс: 1.40-1.77с."
)
add_table(
    ["ТОП-5 быстрых", "Пол", "Время", "ТОП-5 медленных", "Пол", "Время"],
    [
        ["Pulcherrima", "жен.", "1.40с", "Iapetus", "муж.", "1.77с"],
        ["Zubenelgenubi", "муж.", "1.43с", "Kore", "жен.", "1.69с"],
        ["Achernar", "жен.", "1.46с", "Rasalgethi", "муж.", "1.69с"],
        ["Alnilam", "муж.", "1.47с", "Enceladus", "муж.", "1.65с"],
        ["Fenrir", "муж.", "1.48с", "Gacrux", "жен.", "1.63с"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 7. КАЧЕСТВО ЗВУКА
# ═══════════════════════════════════════
add_heading("7. Качество звука: субъективная оценка", 1)

doc.add_paragraph(
    "Субъективное прослушивание экспертом Oki-Toki. Шкала 1-10 (10 = неотличимо от живого оператора)."
)

add_heading("7.1. Оценка естественности", 2)
add_table(
    ["TTS-система", "Тип", "Оценка", "Комментарий"],
    [
        ["Google Chirp3-HD Leda", "API", "9-10 / 10", "Практически неотличимо от живого оператора"],
        ["Google Chirp3-HD Puck", "API", "9-10 / 10", "Мужской голос такого же качества"],
        ["StyleTTS2 UA", "Локальн. GPU", "9-10 / 10*", "Отличное качество (*данные из отдельного теста)"],
        ["ElevenLabs v2", "API", "7-8 / 10", "Хорошее качество, лучше OpenAI для UA"],
        ["Google Wavenet-B", "API", "7-8 / 10", "Хорошо, но ощущается синтез"],
        ["Azure PolinaNeural", "API", "5-6 / 10", "Среднее. Монотонная, без эмоций"],
        ["OpenAI tts-1", "API", "5-6 / 10", "Среднее качество для украинского"],
        ["Edge TTS", "API (бесплатн.)", "3-4 / 10", "Роботизированное, плохие ударения"],
        ["XTTS v2", "Локальн. GPU", "5-6 / 10*", "Клонирование нестабильно (50/50)"],
    ]
)

add_heading("7.2. Эмоциональность", 2)
add_table(
    ["TTS-система", "Различимость эмоций", "Комментарий"],
    [
        ["Google Chirp3-HD", "Чётко различимы ✓", "Автоматически из контекста текста"],
        ["ElevenLabs v2", "Частично", "Некоторые интонационные отличия"],
        ["Все остальные", "Не различимы", "Все эмоции звучат одинаково"],
    ]
)

add_heading("7.3. Качество украинского произношения", 2)
add_table(
    ["Параметр", "Azure Polina", "Wavenet-B", "Chirp3-HD Leda", "ElevenLabs", "OpenAI"],
    [
        ["Ударения", "Хорошо", "Хорошо", "Отлично", "Хорошо", "Средне"],
        ["Артикуляция", "Хорошо", "Средне", "Отлично", "Хорошо", "Средне"],
        ["Паузы", "Средне", "Средне", "Отлично", "Хорошо", "Средне"],
        ["Интонация", "Монотонная", "Слабая", "Живая", "Хорошая", "Средняя"],
    ]
)

add_heading("7.4. Качество после даунсемплинга (→ 8kHz)", 2)
doc.add_paragraph(
    "Все провайдеры генерируют аудио в 16-48kHz, которое конвертируется в 8kHz для телефонии. "
    "Качество после конвертации приемлемо для телефонного канала у всех провайдеров."
)

add_heading("7.5. Рекомендуемый голос", 2)
p = doc.add_paragraph()
run = p.add_run("Выбор эксперта: uk-UA-Chirp3-HD-Leda (женский)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
doc.add_paragraph("Живые, естественные интонации", style='List Bullet')
doc.add_paragraph("Правильные эмоции из контекста текста", style='List Bullet')
doc.add_paragraph("Приятный тембр голоса", style='List Bullet')
doc.add_paragraph("Корректное украинское произношение", style='List Bullet')
doc.add_paragraph("Готов к продакшину для колл-центра", style='List Bullet')

add_heading("7.6. Итоговая оценка качества", 2)
add_table(
    ["Критерий", "Azure", "Wavenet", "Chirp3-HD", "ElevenLabs", "OpenAI"],
    [
        ["Естественность", "5-6", "7-8", "9-10", "7-8", "5-6"],
        ["Эмоциональность", "1", "2", "9", "5", "2"],
        ["Произношение UA", "7", "7", "9", "7", "5"],
        ["Приятность", "6", "7", "9", "7", "5"],
        ["ОБЩИЙ БАЛЛ", "5.0", "5.8", "9.0", "6.5", "4.3"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 8. ОПТИМИЗАЦИЯ
# ═══════════════════════════════════════
add_heading("8. Оптимизация скорости: методы и результаты", 1)

add_heading("8.1. Варианты пайплайна", 2)
add_table(
    ["Вариант пайплайна", "Время", "Эмоции?"],
    [
        ["GPT текст + GPT SSML + Azure TTS", "5.8с", "Нет"],
        ["GPT текст + Structured SSML + Azure", "4.3с", "Нет"],
        ["GPT текст + шаблонный SSML + Azure", "3.3с", "Нет"],
        ["GPT текст + Chirp3-HD (без SSML)", "3.0с", "Да"],
        ["Один GPT (текст+SSML) + Azure", "2.8с", "Нет"],
        ["GPT текст + Chirp3-HD параллельно", "2.3с", "Да"],
        ["Кеш частых фраз + Azure", "1.0с", "Нет"],
    ]
)

add_heading("8.2. Оптимизация Chirp3-HD", 2)
doc.add_paragraph(
    "Параллельная генерация: текст → 2-3 предложения → gRPC-клиенты одновременно → склейка PCM."
)
add_table(
    ["Метод", "Время (150 симв.)", "CPS", "Выигрыш"],
    [
        ["Базовый (1 запрос)", "1.41с", "104", "—"],
        ["gRPC warmup", "1.53с", "95", "+2%"],
        ["Параллель ×2 + warmup", "0.82-1.00с", "146-178", "+36-42%"],
        ["Параллель ×3 + warmup", "0.70-0.87с", "168-209", "+38-44%"],
    ]
)

doc.add_paragraph()
add_bold_text("Стабильный результат (5 запусков, параллель ×2):")
doc.add_paragraph("Среднее: 0.915с | Мин: 0.862с | Макс: 1.046с | CPS: ~160")

add_heading("8.3. Европейский endpoint", 2)
doc.add_paragraph(
    "Chirp3-HD доступен только на серверах US. Европейский endpoint не поддерживает эту модель."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 9. СТОИМОСТЬ
# ═══════════════════════════════════════
add_heading("9. Стоимость провайдеров", 1)

add_table(
    ["TTS-система", "Тип", "Цена за 1M символов", "Бесплатный объём"],
    [
        ["Google Standard", "API", "$4", "1M симв./мес"],
        ["Azure Neural (S0)", "API", "$16", "500K симв./мес (F0)"],
        ["Google Wavenet", "API", "$16", "1M симв./мес"],
        ["OpenAI tts-1", "API", "$15", "Нет"],
        ["Google Chirp3-HD", "API", "Уточнять*", "Есть бесплатный tier"],
        ["ElevenLabs", "API", "~$330**", "10K симв./мес"],
        ["StyleTTS2 UA", "Локальный", "Бесплатно (GPU)", "Требует GPU"],
        ["XTTS v2", "Локальный", "Бесплатно (GPU)", "Требует GPU"],
        ["Edge TTS", "API (бесплатн.)", "Бесплатно", "Без лимитов"],
    ]
)
doc.add_paragraph()
add_note("* Chirp3-HD — ценообразование может отличаться, уточнять в Google Cloud Console")
add_note("** ElevenLabs — $99/мес за 100K символов, $1/10K сверх лимита")

doc.add_page_break()

# ═══════════════════════════════════════
# 10. АРХИТЕКТУРА
# ═══════════════════════════════════════
add_heading("10. Рекомендуемая архитектура для продакшина", 1)

add_heading("10.1. Основной пайплайн", 2)
p = doc.add_paragraph()
run = p.add_run(
    "  Входящий звонок\n"
    "       ↓\n"
    "  Распознавание речи (STT)\n"
    "       ↓\n"
    "  GPT-4o-mini генерирует текст ответа (~1.5с)\n"
    "       ↓\n"
    "  Google Chirp3-HD озвучивает (параллельно, ~0.9с)\n"
    "       ↓\n"
    "  WAV 8kHz → телефонная линия\n"
    "\n"
    "  Общее время: ~2.4с"
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_heading("10.2. Гибридная схема", 2)
add_table(
    ["Тип фразы", "TTS-провайдер", "Время", "Пример"],
    [
        ["Стандартная", "Кеш (предгенерация)", "<0.1с", "«Добрый день, компания Oki-Toki»"],
        ["Динамическая", "Chirp3-HD (параллельно)", "~0.9с", "Ответ на вопрос клиента"],
        ["Fallback", "Azure Neural", "~0.3с", "При ошибке Chirp3-HD"],
    ]
)

add_heading("10.3. Fallback-стратегия", 2)
doc.add_paragraph(
    "При сбое Chirp3-HD — автоматический fallback на Azure PolinaNeural (быстрее, но без эмоций). "
    "При сбое обоих — воспроизведение заранее записанного аудио."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 11. ВЫВОДЫ
# ═══════════════════════════════════════
add_heading("11. Итоговые выводы и рекомендации", 1)

add_heading("11.1. Итоговое сравнение всех 9 TTS", 2)
add_table(
    ["TTS", "Тип", "Скорость UA", "CPS", "Качество", "Эмоции", "Языки"],
    [
        ["Azure Neural", "API", "0.27с", "541", "5-6/10", "Нет", "6/6"],
        ["Google Standard", "API", "0.98с", "149", "4/10", "Нет", "6/6"],
        ["Google Wavenet", "API", "0.75с", "195", "7-8/10", "Нет", "6/6"],
        ["Google Chirp3-HD ⭐", "API", "1.52с*", "96*", "9-10/10", "Да ✓", "6/6"],
        ["ElevenLabs", "API", "1.94с", "75", "7-8/10", "Частично", "6/6"],
        ["OpenAI tts-1", "API", "3.82с", "38", "5-6/10", "Нет", "6/6"],
        ["StyleTTS2 UA", "GPU", "4.86с", "10", "9-10/10", "Нет", "1/6"],
        ["XTTS v2", "GPU", "50с", "14.5", "5-6/10", "Нет", "5/6"],
        ["Edge TTS", "Бесплат.", "5.5с", "~150", "3-4/10", "Нет", "6/6"],
    ]
)
add_note("* Chirp3-HD с параллельной генерацией: 0.86-1.0с / CPS ~160")

add_heading("11.2. Рекомендация", 2)
p = doc.add_paragraph()
run = p.add_run("Основной провайдер: Google Chirp3-HD")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

doc.add_paragraph("Единственный провайдер с реальным эмоциональным синтезом", style='List Bullet')
doc.add_paragraph("30 голосов для украинского языка", style='List Bullet')
doc.add_paragraph("Не требует SSML — экономия 2-4 секунды", style='List Bullet')
doc.add_paragraph("С параллельной генерацией: ~0.9с для 150 символов", style='List Bullet')
doc.add_paragraph("Готов к продакшину (подтверждено экспертной оценкой)", style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Резервный: Microsoft Azure Neural")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x44, 0x88)
doc.add_paragraph("Самый быстрый (0.1-0.3с с warmup)", style='List Bullet')
doc.add_paragraph("Используется как fallback при сбоях Chirp3-HD", style='List Bullet')

add_heading("11.3. Что НЕ рекомендуется", 2)
doc.add_paragraph("SSML через LLM — +2-4с без результата", style='List Bullet')
doc.add_paragraph("OpenAI TTS — медленный для UA (3.8с), среднее качество", style='List Bullet')
doc.add_paragraph("ElevenLabs — хорошее качество, но дорого ($330/1M символов)", style='List Bullet')
doc.add_paragraph("Локальные модели (StyleTTS2, XTTS) — медленные, требуют GPU, ограничение языков", style='List Bullet')
doc.add_paragraph("Edge TTS — роботизированное звучание", style='List Bullet')
doc.add_paragraph("Qwen3-TTS — 171с генерация, неприемлемо", style='List Bullet')

# ═══════════════════════════════════════
# СОХРАНЕНИЕ
# ═══════════════════════════════════════
desktop = os.path.join(os.path.expanduser("~"), "Desktop")
filepath = os.path.join(desktop, "SSML_Отчёт_Oki-Toki.docx")
doc.save(filepath)
print(f"Отчёт сохранён: {filepath}")
print(f"Разделов: 11")
print(f"TTS-провайдеров: 9 (+2 отклонённых)")
