import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ═══════════════════════════════════════
# Стилі
# ═══════════════════════════════════════
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_bold_text(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# ═══════════════════════════════════════
# ТИТУЛЬНА СТОРІНКА
# ═══════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("АНАЛИТИЧЕСКИЙ ОТЧЁТ")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Исследование TTS-провайдеров и SSML-разметки\nдля платформы Oki-Toki")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Формат: WAV 8kHz 16bit mono (телефония)\n6 языков • 4 провайдера • 30+ голосов\nSSML-разметка • Эмоциональный синтез • Оптимизация скорости")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x88)

for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Февраль 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════
# СОДЕРЖАНИЕ
# ═══════════════════════════════════════
add_heading("Содержание", 1)
sections = [
    "1. Цели и задачи исследования",
    "2. Протестированные TTS-провайдеры",
    "3. Сравнение скорости генерации (украинский язык)",
    "4. Мультиязычное тестирование (6 языков)",
    "5. SSML-разметка: тестирование и выводы",
    "6. Эмоциональный синтез: сравнение подходов",
    "7. Качество звука: субъективная оценка",
    "8. Оптимизация скорости: методы и результаты",
    "9. Стоимость провайдеров",
    "10. Рекомендуемая архитектура для продакшина",
    "11. Итоговые выводы и рекомендации",
]
for s in sections:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ═══════════════════════════════════════
# 1. ЦЕЛИ И ЗАДАЧИ
# ═══════════════════════════════════════
add_heading("1. Цели и задачи исследования", 1)
doc.add_paragraph(
    "Цель исследования — определить оптимальный TTS-провайдер для платформы Oki-Toki "
    "(облачный колл-центр) с учётом следующих требований:"
)
requirements = [
    "Формат аудио: WAV 8kHz 16bit mono (стандарт телефонии)",
    "Скорость генерации: менее 1 секунды для фраз до 150 символов",
    "Качество голоса: естественное звучание, максимально близкое к живому оператору",
    "Эмоциональный синтез: способность передавать эмоции (радость, сочувствие, спокойствие)",
    "Мультиязычность: поддержка UA, EN, RU, PL, ES, TR",
    "SSML-разметка: возможность управления просодией (скорость, тон, громкость, паузы)",
    "Стоимость: оптимальное соотношение цена/качество при высоких объёмах",
]
for r in requirements:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# 2. ПРОТЕСТИРОВАННЫЕ ПРОВАЙДЕРЫ
# ═══════════════════════════════════════
add_heading("2. Протестированные TTS-провайдеры", 1)

add_heading("2.1. Microsoft Azure Cognitive Services", 2)
doc.add_paragraph(
    "Тестировались Neural-голоса: uk-UA-PolinaNeural (жен.) и uk-UA-OstapNeural (муж.). "
    "Azure поддерживает WebSocket-соединение с pre-connect и warmup, что обеспечивает "
    "минимальную задержку при повторных запросах. Формат вывода: Raw8Khz16BitMonoPcm."
)
add_table(
    ["Параметр", "Значение"],
    [
        ["Голоса (UA)", "PolinaNeural (жен.), OstapNeural (муж.)"],
        ["Протокол", "WebSocket с pre-connect"],
        ["Warmup", "Поддерживается (открытие соединения + тестовый запрос)"],
        ["Формат", "Raw PCM 8kHz 16bit mono"],
        ["SSML", "Ограниченная поддержка (см. раздел 5)"],
        ["Эмоции", "Не поддерживаются для украинского языка"],
    ]
)

doc.add_paragraph()
add_heading("2.2. Google Cloud Text-to-Speech", 2)

add_heading("2.2.1. Standard (Standard-B)", 3)
doc.add_paragraph(
    "Базовая модель Google. Один женский голос для украинского языка (uk-UA-Standard-B). "
    "Быстрая генерация, но роботизированное звучание."
)

add_heading("2.2.2. WaveNet (Wavenet-B)", 3)
doc.add_paragraph(
    "Нейросетевая модель. Один женский голос (uk-UA-Wavenet-B). "
    "Улучшенное качество по сравнению со Standard, но всё ещё заметно синтетическое звучание. "
    "Поддерживает SSML-управление скоростью (rate) в обоих направлениях."
)

add_heading("2.2.3. Chirp3-HD (30 голосов)", 3)
doc.add_paragraph(
    "Новейшая модель Google (2024-2025). 30 украинских голосов: 14 женских + 16 мужских. "
    "Ключевая особенность — автоматическое распознавание эмоций из контекста текста "
    "без какой-либо SSML-разметки. Наиболее естественное звучание среди всех протестированных."
)
add_table(
    ["Параметр", "Standard-B", "Wavenet-B", "Chirp3-HD"],
    [
        ["Голоса (UA)", "1 жен.", "1 жен.", "14 жен. + 16 муж."],
        ["Качество", "Роботизированное", "Среднее", "Естественное"],
        ["Эмоции", "Нет", "Нет", "Автоматические"],
        ["API", "Batch (synthesize)", "Batch (synthesize)", "Streaming (streaming_synthesize)"],
        ["SSML", "Поддерживает", "Поддерживает rate", "Не требуется"],
    ]
)

doc.add_paragraph()
add_heading("2.3. Не тестировались (теоретический обзор)", 2)
doc.add_paragraph(
    "ElevenLabs и OpenAI TTS не были включены в основное тестирование по следующим причинам:"
)
doc.add_paragraph("ElevenLabs — высокая стоимость ($99/мес за 100K символов), отсутствие украинских голосов для телефонии", style='List Bullet')
doc.add_paragraph("OpenAI TTS — нет нативных украинских голосов, высокая задержка (2-4с), ограниченная настройка формата", style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# 3. СРАВНЕНИЕ СКОРОСТИ (УКРАИНСКИЙ)
# ═══════════════════════════════════════
add_heading("3. Сравнение скорости генерации (украинский язык)", 1)

doc.add_paragraph(
    "Тест проводился на одном и том же тексте длиной ~300 символов. "
    "Azure использовал WebSocket pre-connect с warmup. "
    "Результаты получены на бесплатном тарифе Azure (F0) — на платном (S0) Azure будет стабильнее."
)

add_heading("3.1. Результаты теста (300 символов, украинский)", 2)
add_table(
    ["#", "Модель", "Время", "Комментарий"],
    [
        ["1", "Azure PolinaNeural", "0.27с", "Самый быстрый (с warmup)"],
        ["2", "Google Wavenet-B", "0.75с", "Стабильный результат"],
        ["3", "Google Standard-B", "0.98с", "Быстрый, но роботизированный"],
        ["4", "Azure OstapNeural", "1.29с", "Нестабильный (F0 tier)"],
        ["5", "Google Chirp3-HD Leda", "1.52с", "Лучшее качество"],
        ["6", "Google Chirp3-HD Puck", "1.48с", "Мужской голос"],
        ["7", "Google Chirp3-HD Kore", "1.69с", "Женский голос"],
    ]
)

doc.add_paragraph()
add_heading("3.2. Результаты теста (~150 символов, украинский)", 2)
add_table(
    ["Модель", "Базовый", "С оптимизацией", "Метод оптимизации"],
    [
        ["Azure Polina (warmup)", "0.27с", "0.10-0.30с", "WebSocket pre-connect"],
        ["Google Wavenet-B", "0.75с", "—", "Нет возможности"],
        ["Google Chirp3-HD", "1.41с", "0.86-1.00с", "Параллельная генерация ×2-3"],
    ]
)

doc.add_paragraph()
add_bold_text("Вывод по скорости:")
doc.add_paragraph(
    "Azure — самый быстрый провайдер (0.1-0.3с на платном тарифе). "
    "Chirp3-HD медленнее (1.4-1.9с базовый), но оптимизируется параллельной генерацией до ~0.9с. "
    "Для колл-центра допустимая задержка TTS — до 1 секунды."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 4. МУЛЬТИЯЗЫЧНОЕ ТЕСТИРОВАНИЕ
# ═══════════════════════════════════════
add_heading("4. Мультиязычное тестирование (6 языков)", 1)

doc.add_paragraph(
    "Тест проводился на текстах ~300 символов на каждом языке. "
    "Сравнивались Azure Neural и Google Wavenet — оба поддерживают все 6 языков."
)

add_table(
    ["Язык", "Azure голос", "Azure время", "Google голос", "Google время", "Победитель"],
    [
        ["UA", "PolinaNeural", "0.27с", "Wavenet-B", "0.75с", "Azure"],
        ["EN", "JennyNeural", "0.32с", "Wavenet-F", "0.68с", "Azure"],
        ["RU", "SvetlanaNeural", "0.29с", "Wavenet-A", "0.71с", "Azure"],
        ["PL", "AgnieszkaNeural", "0.35с", "Wavenet-A", "0.82с", "Azure"],
        ["ES", "ElviraNeural", "0.31с", "Wavenet-C", "0.74с", "Azure"],
        ["TR", "EmelNeural", "0.28с", "Wavenet-C", "0.79с", "Azure"],
    ]
)

doc.add_paragraph()
add_bold_text("Вывод:")
doc.add_paragraph(
    "Azure стабильно быстрее Google Wavenet на всех 6 языках (в 2-3 раза). "
    "Однако Azure не поддерживает эмоциональный синтез ни на одном языке. "
    "Google Chirp3-HD доступен не для всех языков — необходимо проверять поддержку."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 5. SSML-РАЗМЕТКА
# ═══════════════════════════════════════
add_heading("5. SSML-разметка: тестирование и выводы", 1)

doc.add_paragraph(
    "SSML (Speech Synthesis Markup Language) — стандарт разметки для управления синтезом речи. "
    "Позволяет задавать скорость, высоту, громкость, паузы и эмфазис. "
    "Мы протестировали SSML-разметку на всех трёх провайдерах."
)

add_heading("5.1. Azure Neural — SSML-поддержка", 2)
add_table(
    ["SSML-параметр", "Результат", "Комментарий"],
    [
        ["rate (замедление)", "Работает", "50-90% — заметный эффект"],
        ["rate (ускорение)", "Не работает", "100-150% — звучит одинаково"],
        ["pitch", "Минимальный эффект", "Меняет тембр, не эмоцию"],
        ["volume", "Не работает", "Никакого эффекта"],
        ["break (паузы)", "Работает", "Единственный надёжный параметр"],
        ["emphasis", "Не работает", "Не поддерживается для Neural"],
        ["express-as (эмоции)", "Не работает", "Не поддерживается для uk-UA"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Важное открытие: ограничения SSML одинаковы для ВСЕХ языков Azure Neural. "
    "Тест на en-US-JennyNeural показал те же результаты, что и uk-UA-PolinaNeural. "
    "Это ограничение модели, а не языка."
)

add_heading("5.2. Google Wavenet — SSML-поддержка", 2)
add_table(
    ["SSML-параметр", "Результат", "Комментарий"],
    [
        ["rate (замедление)", "Работает", "x-slow даёт +99% длительности"],
        ["rate (ускорение)", "Работает", "x-fast даёт -33% длительности"],
        ["pitch", "Минимальный эффект", "Заметно только при экстремальных значениях"],
        ["volume", "Не работает", "Никакого эффекта"],
        ["break (паузы)", "Работает", "Надёжный параметр"],
        ["emphasis", "Минимальный эффект", "Слабо заметно"],
    ]
)

add_heading("5.3. Google Chirp3-HD — SSML не требуется", 2)
doc.add_paragraph(
    "Chirp3-HD автоматически распознаёт эмоциональный контекст из текста и добавляет "
    "соответствующую интонацию. SSML-разметка не поддерживается и не нужна."
)

doc.add_paragraph()
add_heading("5.4. SSML + LLM: тестирование пайплайна", 2)
doc.add_paragraph(
    "Был протестирован пайплайн, где GPT-4o-mini генерирует SSML-разметку для ответа бота:"
)
doc.add_paragraph("Диалог → GPT генерирует текст ответа (~1.5с) → GPT добавляет SSML-разметку (~4.0с) → Azure TTS (~0.3с)", style='List Bullet')
doc.add_paragraph("Общее время: ~5.8с на один ответ", style='List Bullet')
doc.add_paragraph("2 из 10 тестов — невалидный SSML (fallback на plain text)", style='List Bullet')
doc.add_paragraph("Результат: SSML-разметка не давала заметного эмоционального эффекта", style='List Bullet')

doc.add_paragraph()
add_bold_text("Вывод по SSML:")
doc.add_paragraph(
    "SSML-разметка через LLM — это дополнительные 2-4 секунды задержки без ощутимого результата. "
    "Ни Azure, ни Wavenet не способны передать эмоции через SSML-параметры. "
    "Chirp3-HD решает задачу эмоционального синтеза без SSML, экономя время и сложность."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 6. ЭМОЦИОНАЛЬНЫЙ СИНТЕЗ
# ═══════════════════════════════════════
add_heading("6. Эмоциональный синтез: сравнение подходов", 1)

doc.add_paragraph(
    "Проведён сравнительный тест на 3 эмоциях (радость, грусть, спокойствие) "
    "с одинаковыми текстами для каждого провайдера."
)

add_heading("6.1. Тест: Chirp3-HD vs Wavenet (без SSML) vs Wavenet (с SSML)", 2)
add_table(
    ["Эмоция", "Chirp3-HD Leda", "Wavenet (без SSML)", "Wavenet (с SSML)"],
    [
        ["Happy", "Радостная интонация ✓", "Монотонно", "Монотонно"],
        ["Sad", "Сочувственный тон ✓", "Монотонно", "Монотонно"],
        ["Calm", "Спокойный, ровный ✓", "Монотонно", "Монотонно"],
    ]
)

doc.add_paragraph()
add_heading("6.2. Тест: Azure (Polina + Ostap) vs Chirp3-HD (Leda + Puck)", 2)
add_table(
    ["Эмоция", "Azure Polina", "Azure Ostap", "Chirp3 Leda", "Chirp3 Puck"],
    [
        ["Happy", "0.78-1.10с", "0.78с", "1.90с", "1.69с"],
        ["Sad", "1.50с", "0.93с", "1.86с", "1.98с"],
        ["Calm", "1.21с", "0.89с", "1.71с", "2.04с"],
    ]
)
doc.add_paragraph(
    "Azure быстрее, но все эмоции звучат одинаково. "
    "Chirp3-HD медленнее, но каждая эмоция звучит по-разному — радостный текст звучит радостно, "
    "грустный — сочувственно."
)

add_heading("6.3. Все 30 голосов Chirp3-HD", 2)
doc.add_paragraph(
    "Протестированы все 30 украинских голосов Chirp3-HD на одном тексте (146 символов). "
    "Средняя скорость: 1.56с. Разброс: 1.40-1.77с."
)
add_table(
    ["ТОП-5 быстрых", "Пол", "Время", "ТОП-5 медленных", "Пол", "Время"],
    [
        ["Pulcherrima", "жен.", "1.40с", "Iapetus", "муж.", "1.77с"],
        ["Zubenelgenubi", "муж.", "1.43с", "Kore", "жен.", "1.69с"],
        ["Achernar", "жен.", "1.46с", "Rasalgethi", "муж.", "1.69с"],
        ["Alnilam", "муж.", "1.47с", "Enceladus", "муж.", "1.65с"],
        ["Fenrir", "муж.", "1.48с", "Gacrux", "жен.", "1.63с"],
    ]
)
doc.add_paragraph(
    "Разница между голосами минимальна (~0.3с). Выбор голоса следует делать по "
    "субъективному качеству звучания, а не по скорости."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 7. КАЧЕСТВО ЗВУКА
# ═══════════════════════════════════════
add_heading("7. Качество звука: субъективная оценка", 1)

doc.add_paragraph(
    "Проведено субъективное прослушивание всех протестированных TTS-провайдеров "
    "экспертом Oki-Toki. Оценка по шкале от 1 до 10, где 10 — неотличимо от живого оператора."
)

add_heading("7.1. Оценка естественности голоса", 2)
add_table(
    ["Провайдер / голос", "Оценка", "Комментарий эксперта"],
    [
        ["Azure PolinaNeural", "5-6 / 10", "Средний уровень. Монотонная, без эмоций. Звучит как синтезатор — ровный тон вне зависимости от содержания текста"],
        ["Azure OstapNeural", "5-6 / 10", "Аналогично Polina — стабильный, но безэмоциональный"],
        ["Google Wavenet-B", "7-8 / 10", "Хороший уровень. Более живое звучание чем Azure, но всё ещё ощущается синтез"],
        ["Google Chirp3-HD Leda", "9-10 / 10", "Отличный. Практически неотличимо от живого оператора"],
        ["Google Chirp3-HD Puck", "9-10 / 10", "Отличный. Мужской голос такого же высокого качества"],
    ]
)

add_heading("7.2. Оценка эмоциональности", 2)
doc.add_paragraph(
    "Тестировалось на 3 эмоциях: радость (happy), грусть (sad), спокойствие (calm). "
    "Один и тот же текст с соответствующим эмоциональным контекстом."
)
add_table(
    ["Провайдер", "Различимость эмоций", "Комментарий"],
    [
        ["Azure Neural", "Не различимы", "Все 3 эмоции звучат одинаково монотонно"],
        ["Google Wavenet-B (без SSML)", "Не различимы", "Без SSML — одинаковый тон"],
        ["Google Wavenet-B (с SSML)", "Минимально", "SSML меняет скорость, но не эмоцию"],
        ["Google Chirp3-HD", "Чётко различимы ✓", "Радостный текст звучит радостно, грустный — сочувственно, спокойный — ровно и уверенно"],
    ]
)

doc.add_paragraph()
add_bold_text("Экспертная оценка Chirp3-HD:")
doc.add_paragraph(
    "Эмоции чётко различимы на слух. Chirp3-HD автоматически подбирает правильную интонацию "
    "из контекста текста — без какой-либо дополнительной разметки."
)

add_heading("7.3. Качество украинского произношения", 2)
add_table(
    ["Параметр", "Azure Polina", "Google Wavenet-B", "Google Chirp3-HD Leda"],
    [
        ["Правильность ударений", "Хорошо", "Хорошо", "Отлично"],
        ["Чёткость артикуляции", "Хорошо", "Средне", "Отлично"],
        ["Естественность пауз", "Средне", "Средне", "Отлично"],
        ["Интонационные переходы", "Монотонные", "Слабые", "Живые, естественные"],
        ["Проглатывание слов", "Нет", "Нет", "Нет"],
    ]
)

add_heading("7.4. Качество после даунсемплинга (24kHz → 8kHz)", 2)
doc.add_paragraph(
    "Chirp3-HD генерирует аудио в 24kHz, которое конвертируется в 8kHz для телефонии. "
    "Экспертная оценка: качество после конвертации приемлемо для телефонного канала. "
    "Потеря верхних частот не влияет на разборчивость речи и эмоциональную окраску."
)

add_heading("7.5. Качество склейки при параллельной генерации", 2)
doc.add_paragraph(
    "При параллельной генерации (разбиение текста на 2-3 части) между фрагментами "
    "добавляется пауза 150мс. Экспертная оценка: небольшая пауза слышна, "
    "но воспринимается естественно (как пауза между предложениями). "
    "Для колл-центра — приемлемо."
)

add_heading("7.6. Рекомендуемый голос", 2)
p = doc.add_paragraph()
run = p.add_run("Выбор эксперта: uk-UA-Chirp3-HD-Leda (женский)")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

doc.add_paragraph("Живые, естественные интонации", style='List Bullet')
doc.add_paragraph("Правильные эмоции — радость, сочувствие, спокойствие", style='List Bullet')
doc.add_paragraph("Приятный тембр голоса", style='List Bullet')
doc.add_paragraph("Корректное украинское произношение", style='List Bullet')
doc.add_paragraph("Готов к продакшину для колл-центра", style='List Bullet')

add_heading("7.7. Итоговая оценка качества", 2)
add_table(
    ["Критерий", "Azure Neural", "Wavenet-B", "Chirp3-HD Leda"],
    [
        ["Естественность", "5-6 / 10", "7-8 / 10", "9-10 / 10"],
        ["Эмоциональность", "1 / 10", "2 / 10", "9 / 10"],
        ["Произношение UA", "7 / 10", "7 / 10", "9 / 10"],
        ["Приятность голоса", "6 / 10", "7 / 10", "9 / 10"],
        ["8kHz телефония", "Хорошо", "Хорошо", "Хорошо"],
        ["Общий балл", "5.0 / 10", "5.8 / 10", "9.0 / 10"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 8. ОПТИМИЗАЦИЯ СКОРОСТИ
# ═══════════════════════════════════════
add_heading("8. Оптимизация скорости: методы и результаты", 1)

add_heading("8.1. Методы оптимизации пайплайна", 2)
add_table(
    ["Вариант пайплайна", "Время", "Эмоции?"],
    [
        ["GPT текст + GPT SSML + Azure TTS", "5.8с", "Нет"],
        ["GPT текст + GPT Structured SSML + Azure", "4.3с", "Нет"],
        ["GPT текст + шаблонный SSML + Azure", "3.3с", "Нет"],
        ["GPT текст + Chirp3-HD (без SSML)", "3.0с", "Да"],
        ["Один GPT (текст+SSML) + Azure", "2.8с", "Нет"],
        ["GPT текст + Chirp3-HD параллельно", "2.3с", "Да"],
        ["Кеш частых фраз + Azure", "1.0с", "Нет"],
    ]
)

add_heading("8.2. Оптимизация Chirp3-HD", 2)
doc.add_paragraph(
    "Chirp3-HD оптимизируется методом параллельной генерации: текст разбивается на 2-3 предложения, "
    "каждое генерируется отдельным gRPC-клиентом одновременно, затем аудио склеивается."
)
add_table(
    ["Метод", "Время (150 симв.)", "Выигрыш"],
    [
        ["Базовый (1 запрос)", "1.41с", "—"],
        ["gRPC warmup", "1.53с", "+2%"],
        ["Параллель ×2 + warmup", "0.82-1.00с", "+36-42%"],
        ["Параллель ×3 + warmup", "0.70-0.87с", "+38-44%"],
    ]
)

doc.add_paragraph()
add_bold_text("Стабильный результат параллельной генерации (5 запусков):")
doc.add_paragraph("Среднее: 0.915с | Минимум: 0.862с | Максимум: 1.046с")

add_heading("8.3. Европейский endpoint", 2)
doc.add_paragraph(
    "Попытка использовать европейский endpoint (europe-west4) для снижения сетевой задержки "
    "не удалась — Chirp3-HD доступен только на серверах US. "
    "При размещении сервера в США задержка может быть ниже."
)

doc.add_page_break()

# ═══════════════════════════════════════
# 8. СТОИМОСТЬ
# ═══════════════════════════════════════
add_heading("9. Стоимость провайдеров", 1)

doc.add_paragraph(
    "Расчёт стоимости для объёма 1 млн символов в месяц (типичная нагрузка колл-центра)."
)

add_table(
    ["Провайдер", "Модель", "Цена за 1 млн символов", "1 млн симв./мес"],
    [
        ["Azure", "Neural (S0)", "$16", "$16/мес"],
        ["Google", "Standard", "$4", "$4/мес"],
        ["Google", "Wavenet", "$16", "$16/мес"],
        ["Google", "Chirp3-HD", "$—*", "Уточнять"],
        ["ElevenLabs", "Multilingual", "~$330", "$330/мес"],
        ["OpenAI", "TTS-1", "$15", "$15/мес"],
    ]
)
doc.add_paragraph(
    "* Chirp3-HD — новая модель, ценообразование может отличаться. "
    "На момент тестирования использовался бесплатный объём Google Cloud. "
    "Рекомендуется уточнить актуальные цены в Google Cloud Console."
)

doc.add_paragraph()
add_bold_text("Бесплатные объёмы:")
doc.add_paragraph("Azure F0: 500K символов/мес (ограничение 20 запросов/мин)", style='List Bullet')
doc.add_paragraph("Google: первые 1 млн символов Standard / 1 млн Wavenet бесплатно в месяц", style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# 10. АРХИТЕКТУРА
# ═══════════════════════════════════════
add_heading("10. Рекомендуемая архитектура для продакшина", 1)

add_heading("10.1. Основной пайплайн", 2)
doc.add_paragraph("Рекомендуемый пайплайн для Oki-Toki:")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    "  Входящий звонок\n"
    "       ↓\n"
    "  Распознавание речи (STT)\n"
    "       ↓\n"
    "  GPT-4o-mini генерирует текст ответа (~1.5с)\n"
    "       ↓\n"
    "  Google Chirp3-HD озвучивает (параллельно, ~0.9с)\n"
    "       ↓\n"
    "  WAV 8kHz → телефонная линия\n"
    "\n"
    "  Общее время: ~2.4с"
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_heading("10.2. Оптимизации", 2)

add_bold_text("Кеширование частых фраз (~0с TTS):")
doc.add_paragraph("Заранее сгенерировать аудио для стандартных фраз: приветствие, ожидание, прощание, переключение", style='List Bullet')
doc.add_paragraph("Хранить в Redis/файловой системе", style='List Bullet')
doc.add_paragraph("Время отклика: <100мс", style='List Bullet')

doc.add_paragraph()
add_bold_text("Параллельная генерация Chirp3-HD (~0.9с):")
doc.add_paragraph("Разбить текст на 2-3 предложения", style='List Bullet')
doc.add_paragraph("Генерировать каждое отдельным gRPC-клиентом", style='List Bullet')
doc.add_paragraph("Склеить PCM-аудио с паузами 150мс", style='List Bullet')

doc.add_paragraph()
add_bold_text("Стриминг первого чанка:")
doc.add_paragraph("TTFB (Time To First Byte) Chirp3-HD = ~0.2с", style='List Bullet')
doc.add_paragraph("Начинать воспроизведение первого чанка, пока генерируются остальные", style='List Bullet')

add_heading("10.3. Fallback-стратегия", 2)
doc.add_paragraph(
    "При сбое Chirp3-HD — автоматический fallback на Azure PolinaNeural (быстрее, но без эмоций). "
    "При сбое обоих — воспроизведение заранее записанного аудио с просьбой подождать."
)

add_heading("10.4. Гибридная схема", 2)
add_table(
    ["Тип фразы", "TTS-провайдер", "Время", "Пример"],
    [
        ["Стандартная", "Кеш (предгенерация)", "<0.1с", "«Добрый день, компания Oki-Toki»"],
        ["Динамическая", "Chirp3-HD (параллельно)", "~0.9с", "Ответ на вопрос клиента"],
        ["Fallback", "Azure Neural", "~0.3с", "При ошибке Chirp3-HD"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 11. ВЫВОДЫ
# ═══════════════════════════════════════
add_heading("11. Итоговые выводы и рекомендации", 1)

add_heading("11.1. Итоговое сравнение", 2)
add_table(
    ["Критерий", "Azure Neural", "Google Wavenet", "Google Chirp3-HD"],
    [
        ["Скорость", "★★★★★ (0.1-0.3с)", "★★★☆☆ (0.7-1.2с)", "★★★☆☆ (0.9-1.9с)"],
        ["Естественность", "5-6 / 10", "7-8 / 10", "9-10 / 10"],
        ["Эмоциональность", "1 / 10", "2 / 10", "9 / 10"],
        ["Произношение UA", "7 / 10", "7 / 10", "9 / 10"],
        ["SSML", "★★☆☆☆ (только rate↓)", "★★★☆☆ (rate ↑↓)", "Не нужен"],
        ["Голоса (UA)", "2", "1", "30"],
        ["Мультиязычность", "★★★★★", "★★★★★", "★★★★☆"],
        ["Стоимость", "$16/1M симв.", "$16/1M симв.", "Уточнять"],
        ["ОБЩИЙ БАЛЛ", "5.0 / 10", "5.8 / 10", "9.0 / 10"],
    ]
)

add_heading("11.2. Рекомендация", 2)
p = doc.add_paragraph()
run = p.add_run("Основной провайдер: Google Chirp3-HD")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

doc.add_paragraph("Единственный провайдер с реальным эмоциональным синтезом", style='List Bullet')
doc.add_paragraph("30 голосов для украинского языка (выбор тембра)", style='List Bullet')
doc.add_paragraph("Не требует SSML-разметки — экономия 2-4 секунды на каждом запросе", style='List Bullet')
doc.add_paragraph("При параллельной генерации: ~0.9с для 150 символов", style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Резервный провайдер: Microsoft Azure Neural")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x00, 0x44, 0x88)

doc.add_paragraph("Самый быстрый (0.1-0.3с с warmup)", style='List Bullet')
doc.add_paragraph("Стабильный на платном тарифе S0", style='List Bullet')
doc.add_paragraph("Используется как fallback при сбоях Chirp3-HD", style='List Bullet')

add_heading("11.3. Что НЕ рекомендуется", 2)
doc.add_paragraph("SSML-разметка через LLM — дорого по времени (+2-4с) и не даёт результата", style='List Bullet')
doc.add_paragraph("Structured Output для SSML — техническая элегантность без практической пользы", style='List Bullet')
doc.add_paragraph("Google Standard — роботизированное звучание", style='List Bullet')
doc.add_paragraph("ElevenLabs — чрезмерная стоимость для колл-центра", style='List Bullet')

# ═══════════════════════════════════════
# СОХРАНЕНИЕ
# ═══════════════════════════════════════
desktop = os.path.join(os.path.expanduser("~"), "Desktop")
filepath = os.path.join(desktop, "SSML_Отчёт_Oki-Toki.docx")
doc.save(filepath)
print(f"Отчёт сохранён: {filepath}")
print(f"Страниц: ~15-18")
print(f"Разделов: 11")
