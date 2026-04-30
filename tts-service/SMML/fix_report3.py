import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt

filepath = os.path.join(os.path.expanduser("~"), "Desktop", "TTS_Звіт_Окі-Токі.docx")
doc = Document(filepath)

# ============================================================
# 1. Виправити секції 2.6-2.9 (порядок та описи)
# ============================================================
print("=== 1. Виправляю секції 2.6-2.9 ===")

# Видаляємо ВСЕ що стосується 2.6-2.9 (заголовки + описи)
to_remove = []
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if any(x in txt for x in [
        "2.6. StyleTTS2", "2.7. XTTS v2", "2.8. Qwen3-TTS", "2.9. Piper TTS",
        # Описи — різні варіанти тексту
        "Спеціалізована локальна модель для української мови",
        "Локальна модель з клонуванням голосу. Вимагає GPU",
        "Високоякісний TTS з клонуванням. Підтримує тільки 3 мови",
        "Не працює на Windows (помилка espeakbridge)",
        "Лише українська. Відмінна якість звучання",
    ]):
        to_remove.append((i, p))
        print(f"  Видаляю [{i}]: {txt[:70]}")

for _, p in to_remove:
    p._element.getparent().remove(p._element)

print(f"  Видалено {len(to_remove)} параграфів")

# Знаходимо секцію 3 (заголовок "Тести швидкості")
idx_sec3 = None
for i, p in enumerate(doc.paragraphs):
    if "Тести швидкості" in p.text and p.style.name == "Heading 1":
        idx_sec3 = i
        break

if idx_sec3 is None:
    # Спробуємо інший варіант
    for i, p in enumerate(doc.paragraphs):
        if "3.1." in p.text:
            idx_sec3 = i - 1  # параграф перед 3.1
            break

print(f"  Секція 3 знайдена на позиції: {idx_sec3}")
print(f"  Текст: '{doc.paragraphs[idx_sec3].text[:60]}'")

sec3_element = doc.paragraphs[idx_sec3]._element

items = [
    ("Heading 2", "2.6. StyleTTS2 Ukrainian (локальний GPU)"),
    ("Normal", "Спеціалізована локальна модель для української мови. 31 голос. Вимагає GPU. Відмінна якість звучання (⭐⭐⭐⭐⭐), але повільна генерація (4.86с, CPS: 10). Success rate: 98.5%. Підтримує тільки українську мову."),
    ("Heading 2", "2.7. XTTS v2 — Coqui TTS (локальний GPU)"),
    ("Normal", "Локальна модель з клонуванням голосу. Вимагає GPU. Підтримує EN, RU, PL, ES, TR — НЕ підтримує українську. Клонування нестабільне (50/50). Середній час: 50с, CPS: 14.5."),
    ("Heading 2", "2.8. Qwen3-TTS — Alibaba (локальний GPU) ❌"),
    ("Normal", "Високоякісний TTS з клонуванням. Підтримує тільки 3 мови (EN, RU, ES). Середній час генерації 171с (CPS: 5) — неприйнятно повільний для колл-центру."),
    ("Heading 2", "2.9. Piper TTS (локальний CPU)"),
    ("Normal", "Легка локальна модель на CPU. Найшвидша генерація з усіх провайдерів (CPS: 366-517). Підтримує 6 мов. Але якість голосу дуже низька (3/10) — роботний звук, не підходить для колл-центру."),
]

# addprevious вставляє ПЕРЕД елементом → прямий порядок дає правильну послідовність
for style_name, text in items:
    new_p = doc.add_paragraph(text, style=style_name)
    sec3_element.addprevious(new_p._element)

# Перевіримо
print("  Результат:")
for p in doc.paragraphs:
    if any(x in p.text for x in ["2.5.", "2.6.", "2.7.", "2.8.", "2.9.", "Тести швидкості"]):
        print(f"    [{p.style.name}] {p.text[:80]}")
        if "Тести швидкості" in p.text:
            break


# ============================================================
# 2. Оновити таблицю 6 (3.2 Мультимовний тест)
# ============================================================
print("\n=== 2. Оновлюю таблицю 6 (мультимовний тест) ===")

table6 = doc.tables[6]
table6_element = table6._tbl
parent6 = table6_element.getparent()

header = ["Мова", "Azure", "Wavenet", "Chirp3-HD", "ElevenLabs", "OpenAI tts-1", "Piper TTS"]
rows = [
    ["UA", "0.27с / 541", "0.75с / 195", "1.49с / 98",  "1.94с / 75",  "2.35с / 62",  "0.12с / 517"],
    ["EN", "0.32с / 447", "0.68с / 210", "1.10с / 141", "1.38с / 104", "1.94с / 79",  "0.16с / 422"],
    ["RU", "0.29с / 524", "0.71с / 214", "1.68с / 88",  "1.52с / 100", "1.89с / 78",  "0.17с / 353"],
    ["PL", "0.35с / 414", "0.82с / 177", "1.39с / 109", "1.84с / 82",  "2.22с / 68",  "0.14с / 453"],
    ["ES", "0.31с / 468", "0.74с / 196", "1.34с / 120", "1.94с / 82",  "2.74с / 58",  "0.15с / 493"],
    ["TR", "0.28с / 518", "0.79с / 184", "1.34с / 113", "1.88с / 80",  "2.90с / 52",  "0.20с / 366"],
]

new_table = doc.add_table(rows=7, cols=7, style='Table Grid')
all_rows = [header] + rows
for ri, row_data in enumerate(all_rows):
    for ci, cell_text in enumerate(row_data):
        cell = new_table.cell(ri, ci)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

for ci in range(7):
    for run in new_table.cell(0, ci).paragraphs[0].runs:
        run.font.bold = True

table6_element.addprevious(new_table._tbl)
parent6.remove(table6_element)
print("  Таблицю оновлено: 7 колонок (+ Chirp3-HD, ElevenLabs, OpenAI, Piper)")


# ============================================================
# 3. Оновити Piper в таблиці швидкості (таблиця 5)
# ============================================================
print("\n=== 3. Оновлюю Piper в таблиці швидкості ===")

table5 = doc.tables[5]
for ri, row in enumerate(table5.rows):
    if "Piper" in row.cells[0].text:
        row.cells[2].text = "uk_UA-ukrainian_tts"
        row.cells[3].text = "~0.12с*"
        row.cells[4].text = "517*"
        print(f"  Рядок {ri}: Piper → 0.12с / 517 CPS")
        break

# Оновлюємо примітку
for p in doc.paragraphs:
    if "Chirp3-HD — протестований тільки на UA" in p.text:
        p.text = "* Piper — CPS з окремого тесту, час перерахований для ~150 символів. Якість Piper: 3/10 (роботний звук)."
        print("  Примітку оновлено")
        break


# ============================================================
# 4. Додати Piper в таблицю якості (таблиця 8)
# ============================================================
print("\n=== 4. Додаю Piper в таблицю якості ===")

table8 = doc.tables[8]
new_row = table8.add_row()
new_row.cells[0].text = "Piper TTS"
new_row.cells[1].text = "Локальн. CPU"
new_row.cells[2].text = "★☆☆☆☆ (3)"
new_row.cells[3].text = "★★☆☆☆"
for cell in new_row.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
print("  Piper: ★☆☆☆☆ (3/10)")


# ============================================================
# 5. Оновити таблицю 12 (7. Підтримка мов) — всі моделі
# ============================================================
print("\n=== 5. Оновлюю таблицю підтримки мов ===")

table12 = doc.tables[12]
table12_element = table12._tbl
parent12 = table12_element.getparent()

# API провайдери
api_header = ["Мова", "Azure", "G.Standard", "G.Wavenet", "Chirp3-HD", "ElevenLabs", "OpenAI", "Edge TTS"]
api_data = [
    ["UA", "Polina, Ostap", "1 жін.", "1 жін.", "30 (14♀+16♂)", "✓ (multi)", "✓ (auto)", "= Azure"],
    ["EN", "100+",          "30+",    "30+",    "31",            "✓",         "✓",         "= Azure"],
    ["RU", "3",             "5",      "5",      "30+",           "✓",         "✓",         "= Azure"],
    ["PL", "2",             "5",      "5",      "10+",           "✓",         "✓",         "= Azure"],
    ["ES", "10+",           "4",      "4",      "20+",           "✓",         "✓",         "= Azure"],
    ["TR", "2",             "5",      "5",      "10+",           "✓",         "✓",         "= Azure"],
]

# Локальні моделі
local_header = ["Мова", "StyleTTS2 UA", "XTTS v2", "Qwen3-TTS", "Piper TTS"]
local_data = [
    ["UA", "✓ (31 голос)", "✗", "✗", "✓"],
    ["EN", "✗",            "✓", "✓", "✓"],
    ["RU", "✗",            "✓", "✓", "✓"],
    ["PL", "✗",            "✓", "✗", "✓"],
    ["ES", "✗",            "✓", "✓", "✓"],
    ["TR", "✗",            "✓", "✗", "✓"],
]

# Підпис API
api_label = doc.add_paragraph("Хмарні API провайдери:", style="Normal")
for run in api_label.runs:
    run.font.bold = True
    run.font.size = Pt(10)

# Таблиця API
api_table = doc.add_table(rows=7, cols=8, style='Table Grid')
for ri, row_data in enumerate([api_header] + api_data):
    for ci, cell_text in enumerate(row_data):
        cell = api_table.cell(ri, ci)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
for ci in range(8):
    for run in api_table.cell(0, ci).paragraphs[0].runs:
        run.font.bold = True

# Підпис локальні
local_label = doc.add_paragraph("Локальні моделі:", style="Normal")
for run in local_label.runs:
    run.font.bold = True
    run.font.size = Pt(10)

# Таблиця локальних
local_table = doc.add_table(rows=7, cols=5, style='Table Grid')
for ri, row_data in enumerate([local_header] + local_data):
    for ci, cell_text in enumerate(row_data):
        cell = local_table.cell(ri, ci)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
for ci in range(5):
    for run in local_table.cell(0, ci).paragraphs[0].runs:
        run.font.bold = True

# Вставляємо замість старої таблиці
table12_element.addprevious(api_label._element)
api_label._element.addnext(api_table._tbl)
api_table._tbl.addnext(local_label._element)
local_label._element.addnext(local_table._tbl)
parent12.remove(table12_element)
print("  API таблиця: 8 провайдерів")
print("  Локальна таблиця: 4 моделі")
print("  Piper: ✓ для всіх 6 мов")


# ============================================================
# 6. Заголовок секції 7 → Heading 1
# ============================================================
print("\n=== 6. Заголовок секції 7 ===")
for p in doc.paragraphs:
    if "7. Підтримка мов" in p.text and p.style.name == "Normal":
        p.style = doc.styles["Heading 1"]
        print("  Змінено на Heading 1")
        break


# ============================================================
# 7. Підпис під 3.2
# ============================================================
print("\n=== 7. Підпис 3.2 ===")
for p in doc.paragraphs:
    if "Тестове речення: ~150 символів" in p.text:
        p.text = "Формат: час генерації / CPS (Characters Per Second). ~150 символів на мову. Piper — CPS з окремого тесту."
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.italic = True
        print("  Оновлено")
        break


# ============================================================
# ЗБЕРІГАЄМО
# ============================================================
doc.save(filepath)
print(f"\n✅ Звіт збережено: {filepath}")

# Верифікація
doc2 = Document(filepath)
print(f"\nТаблиць: {len(doc2.tables)}")
for ti, t in enumerate(doc2.tables):
    h = [c.text[:14] for c in t.rows[0].cells]
    print(f"  T{ti}: {len(t.rows)}r × {len(t.columns)}c — {h}")

print("\nСекції 2.5 → 3:")
show = False
for p in doc2.paragraphs:
    if "2.5." in p.text:
        show = True
    if show and p.text.strip():
        print(f"  [{p.style.name:10s}] {p.text[:90]}")
    if "3.1." in p.text:
        break

print("\nМультимовна таблиця:")
for row in doc2.tables[6].rows:
    print(f"  {[c.text[:14] for c in row.cells]}")
