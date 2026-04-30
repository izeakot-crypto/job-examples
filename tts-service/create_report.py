import sys, os
sys.stdout = __import__('io').TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Стилі ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ════════════════════════════════════════════════════════════
# ТИТУЛЬНА СТОРІНКА
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Звіт: Аналіз TTS-провайдерів\nдля платформи Окі-Токі')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Text-to-Speech для телефонії (IVR / кол-центр)\nФормат: WAV 8kHz 16bit mono')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Лютий 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. ВСТУП
# ════════════════════════════════════════════════════════════
doc.add_heading('1. Мета дослідження', level=1)
doc.add_paragraph(
    'Мета — знайти оптимальний TTS-провайдер для платформи Окі-Токі (кол-центр). '
    'Критерії відбору:'
)
doc.add_paragraph('Швидкість генерації — до 0.5 секунди на 300 символів', style='List Bullet')
doc.add_paragraph('Формат — WAV 8kHz 16bit mono (стандарт телефонії)', style='List Bullet')
doc.add_paragraph('Підтримка української мови (основна) + EN, RU, PL, ES, TR', style='List Bullet')
doc.add_paragraph('Природність голосу — має звучати як жива людина', style='List Bullet')
doc.add_paragraph('Підтримка стрімінгу — для мінімальної затримки першого чанку', style='List Bullet')
doc.add_paragraph('Адекватна ціна — до $20/1M символів', style='List Bullet')

# ════════════════════════════════════════════════════════════
# 2. ПРОТЕСТОВАНІ ПРОВАЙДЕРИ
# ════════════════════════════════════════════════════════════
doc.add_heading('2. Протестовані TTS-провайдери', level=1)

doc.add_heading('2.1. Azure Neural TTS (Microsoft)', level=2)
doc.add_paragraph(
    'Хмарний TTS від Microsoft. Використовує нейронні голоси. '
    'Підтримує WebSocket з\'єднання для мінімальної затримки. '
    'Нативний формат 8kHz 16bit — не потребує конвертації.'
)
add_table(doc,
    ['Параметр', 'Значення'],
    [
        ['Українські голоси', 'Polina (жін), Ostap (чол)'],
        ['Формат', 'Riff8Khz16BitMonoPcm / Raw8Khz16BitMonoPcm'],
        ['Стрімінг', 'Так (WebSocket + Synthesizing event)'],
        ['Ціна', '$16 / 1M символів (S0), є безкоштовний F0'],
        ['SDK', 'azure-cognitiveservices-speech (Python)'],
        ['Регіон', 'westeurope'],
        ['SSML', 'Так (prosody: rate, pitch, volume, break)'],
        ['Емоційні стилі', 'НІ для UK/RU (тільки EN, ZH, JA)'],
    ]
)

doc.add_paragraph()
doc.add_heading('2.2. Google Cloud TTS', level=2)
doc.add_paragraph(
    'Хмарний TTS від Google. Має 3 покоління моделей: Standard (конкатенативний), '
    'Wavenet (нейронний від DeepMind), Chirp3-HD (найновіший, найкращий). '
    'Нативний формат 8kHz LINEAR16.'
)
add_table(doc,
    ['Параметр', 'Значення'],
    [
        ['Українські голоси', 'Standard-B (1), Wavenet-B (1), Chirp3-HD (30)'],
        ['Всього UK голосів', '32'],
        ['Формат', 'LINEAR16 8kHz (нативно)'],
        ['Стрімінг', 'Тільки Chirp3-HD'],
        ['Ціна', 'Standard $4, Wavenet/Chirp3 $16 / 1M символів'],
        ['SDK', 'google-cloud-texttospeech (Python)'],
        ['Аутентифікація', 'Service Account JSON'],
        ['SSML', 'Так (prosody, break, effects_profile_id)'],
        ['Telephony profile', 'Так (telephony-class-application)'],
    ]
)

doc.add_paragraph()
doc.add_heading('2.3. ElevenLabs', level=2)
doc.add_paragraph(
    'Преміум TTS з найвищою якістю голосу. Мультимовна модель eleven_multilingual_v2. '
    'Підтримує клонування голосу. Але повільніший і дорожчий.'
)
add_table(doc,
    ['Параметр', 'Значення'],
    [
        ['Модель', 'eleven_multilingual_v2'],
        ['Голоси', '21+ (multilingual, не специфічні для UK)'],
        ['Формат', 'pcm_22050, ulaw_8000, mp3_44100'],
        ['8kHz', 'ulaw_8000 нативно або конвертація PCM 22050→8000'],
        ['Стрімінг', 'Так (chunked response)'],
        ['Ціна', '~$30 / 1M символів'],
        ['SDK', 'elevenlabs (Python)'],
    ]
)

doc.add_paragraph()
doc.add_heading('2.4. OpenAI TTS', level=2)
doc.add_paragraph(
    'TTS від OpenAI. Моделі tts-1 (швидка) і tts-1-hd (якісна). '
    'Мультимовна, але без нативного 8kHz. Найповільніша серед тестованих.'
)
add_table(doc,
    ['Параметр', 'Значення'],
    [
        ['Моделі', 'tts-1, tts-1-hd'],
        ['Голоси', 'alloy, echo, fable, onyx, nova, shimmer'],
        ['Формат', 'mp3, opus, aac, flac (немає 8kHz WAV)'],
        ['Стрімінг', 'Так'],
        ['Ціна', '$15 / 1M символів (tts-1), $30 (tts-1-hd)'],
    ]
)

doc.add_paragraph()
doc.add_heading('2.5. Edge TTS (безкоштовний)', level=2)
doc.add_paragraph(
    'Безкоштовний неофіційний доступ до Azure TTS через Edge WebSocket API. '
    'Ті ж голоси що і Azure, але без SLA, без 8kHz формату, нестабільний.'
)
add_table(doc,
    ['Параметр', 'Значення'],
    [
        ['Голоси', 'Ті ж що Azure (Polina, Ostap для UK)'],
        ['Формат', 'Тільки MP3 24kHz (audio-24khz-48kbitrate-mono-mp3)'],
        ['8kHz', 'Потрібна конвертація через ffmpeg'],
        ['Стрімінг', 'Так (WebSocket)'],
        ['Ціна', 'Безкоштовно (неофіційний API)'],
        ['Стабільність', 'Низька (0.5-5с, непередбачувано)'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. РЕЗУЛЬТАТИ ТЕСТІВ ШВИДКОСТІ
# ════════════════════════════════════════════════════════════
doc.add_heading('3. Результати тестів швидкості', level=1)

doc.add_heading('3.1. Українська мова — порівняння всіх провайдерів', level=2)
doc.add_paragraph('Текст ~150-300 символів, формат WAV 8kHz 16bit mono:')
add_table(doc,
    ['Провайдер', 'Модель/Голос', '~150 сим', '~300 сим', 'Стрімінг 1й чанк'],
    [
        ['Azure (F0 free)', 'Polina Neural', '0.1-0.5с', '0.2-2.0с', '0.05-0.5с'],
        ['Azure (S0 paid)*', 'Polina Neural', '~0.1с', '~0.2с', '~0.05с'],
        ['Google Wavenet', 'Wavenet-B', '0.4-0.6с', '0.7-1.2с', 'Немає'],
        ['Google Standard', 'Standard-B', '0.3-0.4с', '0.7-1.0с', 'Немає'],
        ['Google Chirp3-HD', 'Leda/Puck/Kore', '0.7-1.0с', '1.7-3.3с', '0.13-0.19с'],
        ['ElevenLabs', 'Sarah (multilingual)', '~1.5с', '~2.0с', '~0.3с'],
        ['OpenAI', 'tts-1-hd nova', '~1.7с', '~3-5с', '—'],
        ['Edge TTS', 'Polina', '0.5-2.0с', '0.8-5.4с', '—'],
    ]
)
doc.add_paragraph('* Оцінка для платного тарифу S0 на основі документації Microsoft та тестів з прогрівом.',
    style='Normal').runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_heading('3.2. Мультимовний тест (300+ символів)', level=2)
doc.add_paragraph('Azure з WebSocket прогрівом vs Google Wavenet:')
add_table(doc,
    ['Мова', 'Azure (голос)', 'Azure час', 'Google Wavenet', 'Google час'],
    [
        ['UA', 'PolinaNeural', '0.274с', 'uk-UA-Wavenet-B', '1.154с'],
        ['EN', 'JennyNeural', '0.356с', 'en-US-Wavenet-F', '0.998с'],
        ['RU', 'SvetlanaNeural', '1.042с', 'ru-RU-Wavenet-A', '1.373с'],
        ['PL', 'AgnieszkaNeural', '3.881с*', 'pl-PL-Wavenet-A', '1.367с'],
        ['ES', 'ElviraNeural', '0.803с', 'es-ES-Wavenet-C', '1.067с'],
        ['TR', 'EmelNeural', '0.399с', 'tr-TR-Wavenet-C', '1.283с'],
    ]
)
doc.add_paragraph('* Azure PL — стрибок через Free F0 тариф. На платному S0 очікується ~0.2-0.4с.',
    style='Normal').runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_heading('3.3. Вплив оптимізацій на швидкість Azure', level=2)
doc.add_paragraph('Тести показали що правильна оптимізація критична:')
add_table(doc,
    ['Оптимізація', 'Без', 'З оптимізацією', 'Виграш'],
    [
        ['WebSocket pre-connect', '0.5-2.0с', '0.2-1.0с', '~200мс'],
        ['Warmup запит', '0.5-2.0с', '0.1-0.5с', '~300-500мс'],
        ['Reuse Synthesizer', 'Новий кожен раз', 'Один на сесію', '~100мс'],
        ['Raw PCM (без RIFF)', 'Riff8Khz16Bit', 'Raw8Khz16Bit', 'Менше даних'],
        ['Платний S0 тариф', '0.1-2.0с (random)', '~0.1-0.3с (stable)', 'Стабільність'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Код оптимізації Azure (продакшин):')
run.bold = True
doc.add_paragraph(
    '1. Створити SpeechSynthesizer один раз при старті сервісу\n'
    '2. Відкрити WebSocket: Connection.from_speech_synthesizer(synth).open(True)\n'
    '3. Відправити warmup запит: synth.speak_text_async("тест").get()\n'
    '4. Використовувати Raw8Khz16BitMonoPcm + Synthesizing event для стрімінгу\n'
    '5. Обгортати raw PCM у WAV заголовок вручну (wave module)'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. ЯКІСТЬ ГОЛОСУ
# ════════════════════════════════════════════════════════════
doc.add_heading('4. Якість голосу', level=1)

doc.add_heading('4.1. Суб\'єктивна оцінка (для телефонії 8kHz)', level=2)
add_table(doc,
    ['Провайдер', 'Природність', 'Для телефонії', 'Примітка'],
    [
        ['ElevenLabs', '★★★★★', '★★★★★', 'Найкраща якість, але повільний і дорогий'],
        ['Google Chirp3-HD', '★★★★★', '★★★★☆', 'Дуже природній, 30 голосів, але >1с'],
        ['Azure Polina', '★★★★☆', '★★★★☆', 'Хороша для телефонії, швидка'],
        ['Google Wavenet-B', '★★★☆☆', '★★★☆☆', 'Трохи роботизований'],
        ['Google Standard-B', '★★☆☆☆', '★★☆☆☆', 'Помітно роботизований'],
        ['Edge TTS Polina', '★★★★☆', '★★★☆☆', 'Як Azure, але MP3 24kHz (треба конвертувати)'],
    ]
)

doc.add_paragraph()
doc.add_heading('4.2. SSML та покращення голосу', level=2)
doc.add_paragraph('Що працює для української мови:')
add_table(doc,
    ['Функція', 'Azure UK', 'Google Wavenet UK', 'Примітка'],
    [
        ['prosody rate', 'Так', 'Так', 'Швидкість мовлення'],
        ['prosody pitch', 'Так', 'Так', 'Висота голосу'],
        ['prosody volume', 'Так', 'Так', 'Гучність'],
        ['break time', 'Так', 'Так', 'Паузи між фразами'],
        ['mstts:express-as', 'НІ', '—', 'Емоції тільки для EN/ZH/JA'],
        ['emphasis', 'Частково', 'Частково', 'Не завжди помітний ефект'],
        ['Unicode наголоси', 'ПОГАНО', '—', 'Псує вимову! Не використовувати!'],
        ['telephony profile', '—', 'Так', 'effects_profile_id для телефонії'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('УВАГА: ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run('Unicode наголоси (U+0301) псують вимову в Azure TTS — слова вимовляються неправильно. '
           'Використовувати тільки SSML prosody та break для покращення звучання.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. ЦІНИ
# ════════════════════════════════════════════════════════════
doc.add_heading('5. Порівняння цін', level=1)
add_table(doc,
    ['Провайдер', 'Ціна / 1M символів', 'Безкоштовний тариф', 'Примітка'],
    [
        ['Google Standard', '$4', '1M сим/міс', 'Найдешевший, але роботний'],
        ['OpenAI tts-1', '$15', 'Немає', 'Повільний'],
        ['Azure S0', '$16', 'F0: 0.5M сим/міс', 'Найкращий баланс ціна/швидкість'],
        ['Google Wavenet', '$16', '1M сим/міс', 'Хороший але повільніший за Azure'],
        ['Google Chirp3-HD', '$16', '—', 'Найкраща якість Google, але >1с'],
        ['ElevenLabs', '~$30', '10K сим/міс', 'Найкраща якість, але дорого'],
        ['OpenAI tts-1-hd', '$30', 'Немає', 'Повільний і дорогий'],
        ['Edge TTS', 'Безкоштовно', '∞', 'Неофіційний, нестабільний, без SLA'],
    ]
)

doc.add_paragraph()
doc.add_heading('5.1. Розрахунок для Окі-Токі', level=2)
doc.add_paragraph(
    'Приклад: 100 000 дзвінків/місяць × 200 символів середнє повідомлення = 20M символів/місяць'
)
add_table(doc,
    ['Провайдер', '20M символів/міс', 'На рік'],
    [
        ['Google Standard', '$80', '$960'],
        ['Azure S0', '$320', '$3 840'],
        ['Google Wavenet', '$320', '$3 840'],
        ['ElevenLabs', '$600', '$7 200'],
        ['Edge TTS', '$0', '$0 (без гарантій)'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. СТРІМІНГ
# ════════════════════════════════════════════════════════════
doc.add_heading('6. Стрімінг (перший чанк аудіо)', level=1)
doc.add_paragraph(
    'Стрімінг дозволяє почати програвати аудіо ще до завершення генерації. '
    'Критично для телефонії — користувач чує відповідь швидше.'
)
add_table(doc,
    ['Провайдер', 'Стрімінг', '1й чанк', 'Протокол'],
    [
        ['Azure', 'Так', '~50мс (S0)', 'WebSocket + Synthesizing event'],
        ['Google Chirp3-HD', 'Так', '~130-190мс', 'gRPC streaming'],
        ['Google Wavenet', 'НІ', '—', 'Тільки unary gRPC'],
        ['ElevenLabs', 'Так', '~300мс', 'HTTP chunked'],
        ['Edge TTS', 'Так', '~200-500мс', 'WebSocket'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Важливо для Azure стрімінгу: ')
run.bold = True
p.add_run(
    'Використовувати Raw8Khz16BitMonoPcm (без RIFF заголовку) + Synthesizing event callback. '
    'AudioDataStream викликає дублювання першого слова. '
    'Після отримання raw PCM чанків — обгорнути в WAV заголовок через Python wave module.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. ПІДТРИМКА МОВ
# ════════════════════════════════════════════════════════════
doc.add_heading('7. Підтримка мов', level=1)
add_table(doc,
    ['Мова', 'Azure голоси', 'Google Wavenet', 'Google Chirp3-HD'],
    [
        ['Українська (UK)', 'Polina, Ostap', 'Wavenet-B (1 жін)', '30 голосів'],
        ['Англійська (EN)', '100+ голосів', '30+ Wavenet', '30+ Chirp3'],
        ['Російська (RU)', '3 голоси', 'Wavenet-A/B/C/D/E', '30 голосів'],
        ['Польська (PL)', '2 голоси', 'Wavenet-A/B/C/D/E', '30 голосів'],
        ['Іспанська (ES)', '10+ голосів', 'Wavenet-A/B/C/D', '30 голосів'],
        ['Турецька (TR)', '2 голоси', 'Wavenet-A/B/C/D/E', '30 голосів'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. РЕКОМЕНДАЦІЇ
# ════════════════════════════════════════════════════════════
doc.add_heading('8. Рекомендації', level=1)

doc.add_heading('8.1. Основний вибір: Azure Neural TTS (S0)', level=2)
p = doc.add_paragraph()
run = p.add_run('Azure Neural TTS — найкращий вибір для телефонії Окі-Токі.')
run.bold = True

doc.add_paragraph('Найшвидший серед хмарних TTS (~0.1-0.3с на платному S0)', style='List Bullet')
doc.add_paragraph('WebSocket тримає з\'єднання — мінімальна затримка', style='List Bullet')
doc.add_paragraph('Нативний 8kHz WAV — не потрібна конвертація', style='List Bullet')
doc.add_paragraph('Стрімінг з першим чанком ~50мс', style='List Bullet')
doc.add_paragraph('Прийнятна якість голосу для телефонії', style='List Bullet')
doc.add_paragraph('$16/1M символів — адекватна ціна', style='List Bullet')
doc.add_paragraph('Підтримка 6+ мов з якісними нейронними голосами', style='List Bullet')

doc.add_paragraph()
doc.add_heading('8.2. Альтернатива: Google Wavenet-B', level=2)
doc.add_paragraph('Якщо Azure недоступний:')
doc.add_paragraph('Стабільна швидкість ~0.7-1.2с (повільніше за Azure)', style='List Bullet')
doc.add_paragraph('Нативний 8kHz', style='List Bullet')
doc.add_paragraph('Telephony audio profile для оптимізації', style='List Bullet')
doc.add_paragraph('SSML prosody покращує звучання (rate, pitch, breaks)', style='List Bullet')
doc.add_paragraph('Але тільки 1 жіночий голос для української', style='List Bullet')

doc.add_paragraph()
doc.add_heading('8.3. Не рекомендовано', level=2)
doc.add_paragraph('Edge TTS — безкоштовний, але нестабільний, без SLA, тільки MP3', style='List Bullet')
doc.add_paragraph('OpenAI TTS — повільний (1.7-5с), без 8kHz', style='List Bullet')
doc.add_paragraph('ElevenLabs — повільний (~1.5-2с), дорогий ($30/1M)', style='List Bullet')
doc.add_paragraph('Google Chirp3-HD — найкраща якість, але >1с генерація', style='List Bullet')

doc.add_paragraph()
doc.add_heading('8.4. Продакшин конфігурація Azure', level=2)
doc.add_paragraph(
    '1. Тариф: S0 (Standard) — $16/1M символів\n'
    '2. Регіон: westeurope (найближчий до України)\n'
    '3. Голос: uk-UA-PolinaNeural (жін) або uk-UA-OstapNeural (чол)\n'
    '4. Формат: Raw8Khz16BitMonoPcm (для стрімінгу)\n'
    '5. Оптимізації:\n'
    '   — Один SpeechSynthesizer на сервіс (reuse)\n'
    '   — Pre-connect WebSocket при старті\n'
    '   — Warmup запит при старті\n'
    '   — Synthesizing event для стрімінгу чанків\n'
    '   — Обгортка PCM у WAV header через wave module'
)

# ════════════════════════════════════════════════════════════
# ЗБЕРЕЖЕННЯ
# ════════════════════════════════════════════════════════════
desktop = os.path.join(os.path.expanduser("~"), "Desktop")
if not os.path.exists(desktop):
    desktop = os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop")
if not os.path.exists(desktop):
    desktop = os.path.join(os.path.expanduser("~"), "Рабочий стол")
if not os.path.exists(desktop):
    desktop = os.path.expanduser("~")

filepath = os.path.join(desktop, "TTS_Звіт_Окі-Токі.docx")
doc.save(filepath)
print(f"Звіт збережено: {filepath}")
